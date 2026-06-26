"""
gerar_memorial.py — Gera o memorial final de um relatório preventivo.
Uso: python3 gerar_memorial.py <input.docx> <output.docx>

Parser adaptativo: lê diretamente do .docx, detecta tipo de cada tabela
automaticamente pelo cabeçalho real, sem conversão para markdown.
"""
import sys, os, io, re
sys.path.insert(0, os.path.dirname(__file__))

from parse_report import parse_docx, consolidate, fmt

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Cores ──────────────────────────────────────────────────────────────────────
GRAY_DARK = '808080'   # Itens / código / descrição
GRAY_MID  = 'AEAAAA'  # linha de Total
BLACK     = '000000'

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ── Larguras de coluna em twips (1 cm ≈ 567 twips, total = 10540) ─────────────
# Foto fixa em 1500 em todos os tipos. Restante distribuído proporcionalmente.
# full (6 col):        Foto | Comp | Alt | Desc | Sub | Total
COL_FULL       = [1500, 1800, 1800, 1800, 1800, 1840]
# nodiscount (4 col):  Foto | Comp | Alt | Total
COL_NODISCOUNT = [1500, 3013, 3013, 3014]
# unit (3 col):        Foto | Qty  | Unidade
COL_UNIT       = [1500, 4520, 4520]


# ── Helpers XML ────────────────────────────────────────────────────────────────
def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.upper())
    tcPr.append(shd)


def set_col_width(cell, w_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def write_cell(cell, text, bold=False, align=CENTER, shade=None, size_pt=11):
    if shade:
        set_shading(cell, shade)
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(str(text) if text is not None else '')
    run.font.name = 'Calibri'
    run.font.bold = bold
    run.font.size = Pt(size_pt)


def apply_col_widths(table, col_widths):
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(col_widths):
                set_col_width(cell, col_widths[ci])


# ── Construção do memorial ─────────────────────────────────────────────────────
def add_memorial_table(doc, item):
    htype       = item['htype']
    code        = item['code']
    desc        = item['desc']
    rows        = item['rows']
    total       = item['total']
    unit        = item['unit']
    total_label = item['total_label']
    last_header = item['last_header']

    if htype == 'full':
        ncols = 6
        cw    = COL_FULL
        hdrs  = ['Foto', 'Comp. (m)', 'Altura (m)', 'Desconto (m²)', 'Subtotal',
                 last_header or 'Total (m²)']
    elif htype == 'nodiscount':
        ncols = 4
        cw    = COL_NODISCOUNT
        hdrs  = ['Foto', 'Comp. (m)', 'Altura (m)', last_header or 'Total (m²)']
    else:  # unit
        ncols = 3
        cw    = COL_UNIT
        hdrs  = ['Foto', 'Quantitativo', 'Unidade']

    nrows = 3 + len(rows) + 1   # Itens + código/desc + header + dados + total
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Row 0: "Itens" mesclado, cinza escuro ──
    r0 = table.rows[0]
    r0.cells[0].merge(r0.cells[-1])
    write_cell(r0.cells[0], 'Itens', bold=True, align=CENTER, shade=GRAY_DARK)
    set_col_width(r0.cells[0], sum(cw))

    # ── Row 1: código | descrição (cols 1..n mescladas), cinza escuro ──
    r1 = table.rows[1]
    write_cell(r1.cells[0], code, bold=True, shade=GRAY_DARK)
    set_col_width(r1.cells[0], cw[0])
    r1.cells[1].merge(r1.cells[-1])
    write_cell(r1.cells[1], desc, bold=True, shade=GRAY_DARK)

    # ── Row 2: cabeçalhos das colunas ──
    r2 = table.rows[2]
    for ci, hdr in enumerate(hdrs):
        write_cell(r2.cells[ci], hdr, bold=True, align=CENTER)
        set_col_width(r2.cells[ci], cw[ci])

    # ── Linhas de dados ──
    for ri, r in enumerate(rows):
        tr = table.rows[3 + ri]
        if htype == 'full':
            write_cell(tr.cells[0], r.get('foto', ''), align=CENTER)
            write_cell(tr.cells[1], r.get('comp', ''), align=CENTER)
            write_cell(tr.cells[2], r.get('alt',  ''), align=CENTER)
            write_cell(tr.cells[3], r.get('desc', ''), align=CENTER)
            write_cell(tr.cells[4], r.get('sub',  ''), align=CENTER)
            write_cell(tr.cells[5], r.get('tot',  ''), align=CENTER)
        elif htype == 'nodiscount':
            write_cell(tr.cells[0], r.get('foto', ''), align=CENTER)
            write_cell(tr.cells[1], r.get('comp', ''), align=CENTER)
            write_cell(tr.cells[2], r.get('alt',  ''), align=CENTER)
            write_cell(tr.cells[3], r.get('tot',  ''), align=CENTER)
        else:  # unit
            write_cell(tr.cells[0], r.get('foto', ''), align=CENTER)
            write_cell(tr.cells[1], r.get('qty',  ''), align=CENTER)
            write_cell(tr.cells[2], r.get('unit', ''), align=CENTER)
        for ci in range(ncols):
            set_col_width(tr.cells[ci], cw[ci])

    # ── Linha de Total, cinza médio ──
    tr_tot = table.rows[-1]
    tr_tot.cells[0].merge(tr_tot.cells[-2])
    # Normalizar label: remover fator (X3, multiplicado/dividido por N) — já aplicado nos valores
    clean_label = total_label.strip()
    if '(' in clean_label or 'multiplicado' in clean_label.lower() or 'dividido' in clean_label.lower():
        clean_label = 'Total'
    write_cell(tr_tot.cells[0], clean_label, bold=True, align=CENTER, shade=GRAY_MID)
    val_str = f'{fmt(total)} {unit}' if (htype == 'unit' and unit) else fmt(total)
    write_cell(tr_tot.cells[-1], val_str, bold=True, align=CENTER, shade=GRAY_MID)

    apply_col_widths(table, cw)


# ── Main ───────────────────────────────────────────────────────────────────────
def main(input_path, output_path):
    occurrences  = parse_docx(input_path)
    consolidated = consolidate(occurrences)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin  = section.bottom_margin = Cm(2)

    for idx, (code, item) in enumerate(consolidated.items()):
        add_memorial_table(doc, item)
        if idx < len(consolidated) - 1:
            doc.add_paragraph('')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with open(output_path, 'wb') as f:
        f.write(buf.read())

    print(f'OK Memorial salvo: {output_path}')
    print(f'   {len(consolidated)} itens unicos consolidados')
    for code, item in consolidated.items():
        unit_str = item['unit'] if item['unit'] else item['htype']
        print(f'   {code}: {fmt(item["total"])} {unit_str} ({item["htype"]})')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Uso: python3 gerar_memorial.py <input.docx> <output.docx>')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2])
