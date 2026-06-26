import sys
import os
import re
from docx import Document

sys.path.insert(0, os.path.dirname(__file__))
from parse_report import parse_docx

def expand_range(ref):
    nums = re.findall(r'\d+', ref)
    if not nums:
        return set()
    if len(nums) == 1:
        return {int(nums[0])}
    if ' a ' in ref.lower():
        return set(range(int(nums[0]), int(nums[-1]) + 1))
    return {int(n) for n in nums}

def main(docx_path):
    doc = Document(docx_path)
    body_lines = [p.text for p in doc.paragraphs]
    
    # Extrair texto de tabelas que NÃO são tabelas de cálculo (legendas colocadas por script em tabelas pequenas)
    for t in doc.tables:
        is_calc_table = False
        if len(t.rows) >= 4:
            if t.rows[0].cells and 'Itens' in t.rows[0].cells[0].text:
                is_calc_table = True
                
        if not is_calc_table:
            seen_cells = set()
            for r in t.rows:
                for c in r.cells:
                    if c._tc in seen_cells:
                        continue
                    seen_cells.add(c._tc)
                    for p in c.paragraphs:
                        if p.text.strip():
                            body_lines.append(p.text)
    
    all_fotos = set()
    for text in body_lines:
        for m in re.finditer(r'Foto\s+(\d+)', text):
            all_fotos.add(int(m.group(1)))

    if not all_fotos:
        print('[!] Nenhuma legenda de foto encontrada no corpo do relatorio.')
        return

    max_foto = max(all_fotos)
    expected = set(range(1, max_foto + 1))
    missing  = expected - all_fotos

    standalone = {}
    for i, line in enumerate(body_lines, 1):
        if re.match(r'^\s*Foto\s+\d+\s*$', line):
            n = int(re.search(r'\d+', line).group())
            standalone.setdefault(n, []).append(i)
    duplicates = {n: ls for n, ls in standalone.items() if len(ls) > 1}

    table_fotos = set()
    range_issues = []
    
    occurrences = parse_docx(docx_path)
    for occ in occurrences:
        for row in occ['rows']:
            foto_text = row.get('foto', '')
            if not foto_text:
                continue
            for ref in str(foto_text).split('\n'):
                ref = ref.strip()
                if not ref: continue
                f_ref = 'Foto ' + ref if 'Foto' not in ref else ref
                fotos = expand_range(f_ref)
                table_fotos.update(fotos)
                
                if ' a ' in ref.lower():
                    m = re.search(r'(\d+)\s*a\s*(\d+)', ref.lower())
                    if m:
                        start, end = int(m.group(1)), int(m.group(2))
                        missing_in_range = [n for n in range(start, end + 1) if n not in all_fotos]
                        if missing_in_range:
                            range_issues.append((start, end, missing_in_range))

    table_fotos.discard(0)
    not_captioned = table_fotos - all_fotos

    print('=' * 65)
    print('VERIFICACAO DE LEGENDAS DE FOTOS (V2 - INCLUINDO TABELAS)')
    print('=' * 65)
    print(f'Total de fotos unicas: {len(all_fotos)} (Foto 1 a Foto {max_foto})')

    if missing:
        print(f'\n[!] Lacunas na sequencia ({len(missing)} fotos):')
        print(f'   {sorted(missing)}')
    else:
        print(f'[OK] Sequencia completa de 1 a {max_foto} - sem lacunas')

    if duplicates:
        print(f'\n[X] Legendas DUPLICADAS ({len(duplicates)} fotos):')
        for n, ls in sorted(duplicates.items()):
            print(f'   Foto {n} aparece duplicada')
    else:
        print('[OK] Nenhuma legenda duplicada')

    if not_captioned:
        print(f'\n[X] Fotos referenciadas em tabelas MAS sem legenda ({len(not_captioned)}):')
        print(f'   {sorted(not_captioned)}')
    else:
        print('[OK] Todas as fotos de tabelas possuem legenda no documento')

    if range_issues:
        print(f'\n[X] Intervalos com fotos faltando:')
        for start, end, miss in range_issues:
            print(f'   Foto {start} a {end}: faltam {miss}')
    else:
        print('[OK] Todos os intervalos (Foto X a Y) estao cobertos')

    print('=' * 65)
    total_issues = len(missing) + len(duplicates) + len(not_captioned) + len(range_issues)
    if total_issues == 0:
        print('[OK] Legendas OK - nenhum problema encontrado.')
    else:
        print(f'[X] {total_issues} tipo(s) de problema encontrado(s). Corrija antes de finalizar.')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Uso: python3 verificar_legendas_v2.py <input.docx>')
        sys.exit(1)
    main(sys.argv[1])
