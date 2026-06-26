"""
parse_report.py — utilitários de parsing para a skill relatorio-preventivo.

Lê DIRETAMENTE do .docx via python-docx (sem conversão para markdown).
Adaptativo: detecta tipo de tabela pelo cabeçalho real, usa índice posicional
raw sem deduplicação de células mescladas — isso preserva valores iguais em
colunas adjacentes (ex: Comp=1,00 e Alt=1,00 na mesma linha).
"""
import re
from collections import OrderedDict
from docx import Document


def fmt(v):
    return f"{v:.2f}".replace('.', ',')


def parse_float(s):
    if not s:
        return 0.0
    s = str(s).strip().replace(',', '.')
    s = re.sub(r'[^\d.\-]', '', s)
    try:
        return float(s)
    except Exception:
        return 0.0


def get_cell_text(cell):
    """Texto completo de uma célula (todas as linhas de parágrafo)."""
    return '\n'.join(p.text for p in cell.paragraphs).strip()


def raw_row(row):
    """
    Retorna lista com o texto de CADA célula física pelo índice real.
    NÃO deduplica células mescladas — preserva valores iguais em colunas distintas.
    """
    return [get_cell_text(c) for c in row.cells]


def detect_htype(headers):
    """
    Detecta o tipo da tabela a partir dos headers reais da linha de cabeçalho.
    Retorna: 'full' | 'nodiscount' | 'unit'

    Lógica adaptativa:
      - 'unit'        → header contém 'quantitativo' ou 'unidade'
      - 'full'        → header contém 'desconto' OU 5+ colunas (subtotal implícito)
      - 'nodiscount'  → demais casos (Foto | Comp | Alt | Total)
    """
    h = ' '.join(str(x) for x in headers).lower()
    if 'quantitativo' in h or 'unidade' in h:
        return 'unit'
    if 'desconto' in h or len(headers) >= 5:
        # 5+ colunas sem ser unit → full (subtotal/desconto implícito)
        return 'full'
    return 'nodiscount'


def extract_factor(total_label):
    """
    Extrai o fator multiplicador/divisor do label da linha de Total.
    Ex: 'Total (X3)' → 3.0 (multiplica)
        'Total multiplicado por 25' → 25.0 (multiplica)
        'Total dividido por 10' → 0.1 (multiplica por 1/10)
        'Total dividido por 2' → 0.5
    Retorna 1.0 se não houver fator.
    """
    label = total_label.lower().strip()
    
    # 'Total (X3)' ou 'Total (x3)'
    m = re.search(r'\(x\s*(\d+)\)', label)
    if m:
        return float(m.group(1))
    
    # 'Total multiplicado por 25'
    m = re.search(r'multiplicado\s+por\s+(\d+)', label)
    if m:
        return float(m.group(1))
    
    # 'Total dividido por 10'
    m = re.search(r'dividido\\s+por\\s+(\\d+)', label)
    if m:
        return 1.0 / float(m.group(1))
    
    # 'Total dividido por' (sem número — retorna -1 pra auto-detectar)
    if 'dividido' in label:
        return -1.0
    
    return 1.0


def apply_factor_to_rows(data_rows, factor, htype, declared_total=None):
    """
    Aplica o fator multiplicador/divisor ao campo 'tot' de cada data row.
    Para full/nodiscount: ajusta 'tot'
    Para unit: ajusta 'qty'
    
    Se factor == -1.0: auto-detectar via ratio entre soma dos tots e declared.
    """
    if factor == 1.0:
        # Preencher desconto vazio com 0,00
        if htype == 'full':
            for r in data_rows:
                if 'desc' in r and not r['desc'].strip():
                    r['desc'] = '0,00'
        return
    
    if factor == -1.0 and declared_total and declared_total > 0:
        # Auto-detectar dividido: sum(tots) / declared = divisor
        sum_tots = sum(parse_float(r.get('tot', '0')) for r in data_rows)
        if sum_tots > 0:
            factor = declared_total / sum_tots
        else:
            factor = 1.0
    
    if factor == -1.0:  # couldn't auto-detect
        factor = 1.0
    
    for r in data_rows:
        if htype == 'unit':
            if r.get('qty'):
                orig = parse_float(r['qty'])
                r['qty'] = fmt(orig * factor)
        else:
            if r.get('tot'):
                orig = parse_float(r['tot'])
                r['tot'] = fmt(orig * factor)
            # Ajustar também desconto (desconto vazio → 0,00)
            if 'desc' in r and not r['desc'].strip():
                r['desc'] = '0,00'


def parse_docx(input_path):
    """
    Lê todas as tabelas de dados do .docx e retorna lista de ocorrências.

    Cada ocorrência é um dict:
      code, desc, htype, headers, last_header,
      rows, total_value, total_label, unit
    """
    doc = Document(input_path)
    occurrences = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 4:
            continue

        r0 = raw_row(rows[0])
        r1 = raw_row(rows[1])
        r2 = raw_row(rows[2])

        # Linha 0: deve conter 'Itens' (célula mesclada de título)
        if not r0 or 'Itens' not in r0[0]:
            continue

        # Linha 1: [código, descrição, descrição, ...] (descrição mesclada)
        code = r1[0].strip()
        desc = r1[1].strip() if len(r1) > 1 else ''
        if not code or not re.match(r'^\d+\.\d+', code):
            continue

        # Linha 2: cabeçalhos reais das colunas
        headers = r2
        htype = detect_htype(headers)
        last_header = headers[-1] if headers else 'Total (m²)'

        # Linhas de dados: rows[3] até rows[-2]
        data_rows = []
        unit = ''

        for ri in range(3, len(rows) - 1):
            r = raw_row(rows[ri])
            if not r or not r[0].strip():
                continue
            # Pular linhas de Total intermediárias
            if r[0].strip().lower().startswith('total'):
                continue

            if htype == 'full':
                # Determina número real de colunas da tabela (5 ou 6)
                # 5 cols: Foto(0) | Comp(1) | Alt(2) | Subtotal(3) | Total(4)
                # 6 cols: Foto(0) | Comp(1) | Alt(2) | Desc(3) | Sub(4) | Tot(5)
                ncols = len(headers)
                if ncols >= 6:
                    data_rows.append({
                        'foto': r[0] if len(r) > 0 else '',
                        'comp': r[1] if len(r) > 1 else '',
                        'alt':  r[2] if len(r) > 2 else '',
                        'desc': r[3] if len(r) > 3 else '',
                        'sub':  r[4] if len(r) > 4 else '',
                        'tot':  r[5] if len(r) > 5 else '',
                    })
                else:  # 5 colunas
                    data_rows.append({
                        'foto': r[0] if len(r) > 0 else '',
                        'comp': r[1] if len(r) > 1 else '',
                        'alt':  r[2] if len(r) > 2 else '',
                        'sub':  r[3] if len(r) > 3 else '',
                        'desc': '',
                        'tot':  r[4] if len(r) > 4 else '',
                    })
            elif htype == 'nodiscount':
                # Índices: Foto(0) | Comp(1) | Alt(2) | Tot(3)
                data_rows.append({
                    'foto': r[0] if len(r) > 0 else '',
                    'comp': r[1] if len(r) > 1 else '',
                    'alt':  r[2] if len(r) > 2 else '',
                    'tot':  r[3] if len(r) > 3 else '',
                })
            else:  # unit
                # Índices: Foto(0) | Qty(1) | Unit(2)
                un = r[2] if len(r) > 2 else ''
                if un:
                    unit = un
                data_rows.append({
                    'foto': r[0] if len(r) > 0 else '',
                    'qty':  r[1] if len(r) > 1 else '',
                    'unit': un,
                })

        # Linha de total (última linha da tabela)
        total_row = raw_row(rows[-1])
        total_label = total_row[0].strip() if total_row else 'Total'
        total_value = 0.0

        if htype == 'unit':
            # Para unit: qty em [-2], unidade em [-1]
            if len(total_row) >= 3:
                total_value = parse_float(total_row[-2])
                unit = total_row[-1] or unit
            elif len(total_row) == 2:
                total_value = parse_float(total_row[-1])
            elif len(total_row) == 1:
                # Célula mesclada: "Total 12,00 M/MÊS" — extrair valor+unidade do texto
                combined = total_row[0]
                # Remove label "Total" ou similar no início
                for prefix in ('total', 'subtotal', 'valor'):
                    if combined.lower().startswith(prefix):
                        combined = combined[len(prefix):].strip()
                        break
                # Tenta extrair par "valor unidade" do que sobrou
                m = re.match(r'([\d.,]+)\s*(.*)', combined)
                if m:
                    total_value = parse_float(m.group(1))
                    unit = m.group(2).strip() or unit
            # Fallback: unit das linhas de dados
            if not unit:
                for r in data_rows:
                    if r.get('unit'):
                        unit = r['unit']
                        break
        else:
            # Para full/nodiscount: valor na última célula física
            total_value = parse_float(total_row[-1]) if total_row else 0.0

        # ── Aplicar fator multiplicador/divisor às linhas de dados ──
        factor = extract_factor(total_label)
        apply_factor_to_rows(data_rows, factor, htype, declared_total=total_value)

        occurrences.append({
            'code':        code,
            'desc':        desc,
            'htype':       htype,
            'headers':     headers,
            'last_header': last_header,
            'rows':        data_rows,
            'total_value': total_value,
            'total_label': total_label,
            'unit':        unit,
        })

    return occurrences


def consolidate(occurrences):
    """
    Agrupa ocorrências por código, acumulando rows e somando totais.
    Mantém ordem de primeira aparição.
    """
    consolidated = OrderedDict()

    for occ in occurrences:
        c = occ['code']
        if c not in consolidated:
            consolidated[c] = {
                'code':        c,
                'desc':        occ['desc'],
                'htype':       occ['htype'],
                'last_header': occ['last_header'],
                'rows':        [],
                'total':       0.0,
                'unit':        occ['unit'],
                'total_label': 'Total',
                'labels':      set(),
            }
        consolidated[c]['rows'].extend(occ['rows'])
        consolidated[c]['total'] += occ['total_value']
        consolidated[c]['labels'].add(occ['total_label'])
        if occ['unit'] and not consolidated[c]['unit']:
            consolidated[c]['unit'] = occ['unit']

    # Resolver o label final do Total
    for c, d in consolidated.items():
        lbls = d['labels']
        if len(lbls) == 1:
            d['total_label'] = list(lbls)[0]
        elif any('dividido' in l.lower() for l in lbls):
            d['total_label'] = next(l for l in lbls if 'dividido' in l.lower())
        else:
            d['total_label'] = 'Total'

    return consolidated
