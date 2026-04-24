"""
gerar_memorial.py — Etapa 2: gera o memorial final de um relatório preventivo.
Uso: python3 gerar_memorial.py <input.md> <output.docx>
Após gerar, execute: python3 fix_xml.py <output.docx>
"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from parse_report import (parse_calc_occurrences, parse_itens_tables,
                           consolidate_occurrences, find_memorial_start, fmt)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Cores padrão (não alterar) ────────────────────────────────────────────────
GRAY_DARK = '808080'   # linha de título e cabeçalho "Itens"
GRAY_MID  = 'AEAAAA'  # Subtotal e Total
WHITE     = 'FFFFFF'   # célula de desconto no Subtotal
BLACK     = '000000'

# ── Larguras padrão em cm ─────────────────────────────────────────────────────
WIDTHS = {
    'full':       [7.0, 2.0, 2.0, 2.5, 2.5],  # 5 cols com DESCONTO
    'nodiscount': [8.5, 2.3, 2.3, 2.9],        # 4 cols sem DESCONTO
    'simple':     [12.5, 3.5],                  # 2 cols por unidade
    'itens':      [1.5, 11.0, 2.0, 1.5],        # tabela Itens
}

CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT   = WD_ALIGN_PARAGRAPH.LEFT


# ── Helpers XML ───────────────────────────────────────────────────────────────
def set_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(f'{{{qn("w:shd").split("}")[0][1:]}}}shd'):
        tcPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    va = tcPr.find(f'{{{qn("w:vAlign").split("}")[0][1:]}}}vAlign')
    if va is not None:
        tcPr.insert(list(tcPr).index(va), shd)
    else:
        tcPr.append(shd)


def set_borders(cell, top='single', bottom='single', left='single', right='single',
                right_color='auto'):
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(f'{{{W_NS}}}tcBorders'):
        tcPr.remove(e)
    tcB = OxmlElement('w:tcBorders')

    def add(name, val, color='auto'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), val)
        if val != 'nil':
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
        tcB.append(b)

    add('top', top)
    add('left', left)
    add('bottom', bottom)
    add('right', right, right_color)
    shd_el = tcPr.find(f'{{{W_NS}}}shd')
    if shd_el is not None:
        tcPr.insert(list(tcPr).index(shd_el), tcB)
    else:
        tcPr.append(tcB)


def set_vAlign(cell, val='center'):
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(f'{{{W_NS}}}vAlign'):
        tcPr.remove(e)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)


def cp(cell, text, bold=False, color=BLACK, size=9, align=CENTER, font='Calibri'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    set_vAlign(cell)


def set_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for e in tblPr.findall(f'{{{W_NS}}}tblW'):
        tblPr.remove(e)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(sum(widths_cm) * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblStyle = tblPr.find(f'{{{W_NS}}}tblStyle')
    idx = list(tblPr).index(tblStyle) + 1 if tblStyle is not None else 0
    tblPr.insert(idx, tblW)


def set_row_h(row, twips):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(twips))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)


def gap_para(doc, before_twips=0, after_twips=120):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_twips / 20)
    p.paragraph_format.space_after = Pt(after_twips / 20)
    return p


# ── Construção das tabelas ────────────────────────────────────────────────────
def add_calc_table(doc, d, items_info):
    htype     = d['htype']
    rows_data = d['rows']
    total_val = d['total_sum']
    final_lbl = d['final_label']
    title_txt = d['title']
    tu        = d['total_unit']
    c1        = d['col1_label']
    unit_str  = items_info.get(d['code'], {}).get('unit', tu)

    widths = WIDTHS[htype]
    ncols  = len(widths)
    has_sub = htype in ('full', 'nodiscount')
    nrows = 1 + 1 + len(rows_data) + (1 if has_sub else 0) + 1

    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    ri = 0

    # Título (cinza escuro, texto preto bold, centralizado, colunas mescladas)
    trow = tbl.rows[ri]; ri += 1
    for ci in range(1, ncols):
        trow.cells[0].merge(trow.cells[ci])
    cp(trow.cells[0], title_txt, bold=True, color=BLACK, font='Calibri', align=CENTER)
    set_shd(trow.cells[0], GRAY_DARK)
    set_borders(trow.cells[0], top='single', bottom='single', left='single', right='nil')
    set_row_h(trow, 280)

    # Cabeçalho (sem fundo, texto preto bold)
    hrow = tbl.rows[ri]; ri += 1
    if htype == 'full':
        hdrs = ['REFERÊNCIA', c1, 'ALTURA (m)', 'DESCONTO', f'TOTAL ({tu})']
    elif htype == 'nodiscount':
        hdrs = ['REFERÊNCIA', c1, 'ALTURA (m)', f'TOTAL ({tu})']
    else:
        hdrs = ['REFERÊNCIA', f'TOTAL ({unit_str})']
    for ci, h in enumerate(hdrs):
        cp(hrow.cells[ci], h, bold=True, color=BLACK, font='Calibri', align=CENTER)
        set_borders(hrow.cells[ci], top='nil', bottom='single',
                    left='single' if ci == 0 else 'nil', right='single')
    set_row_h(hrow, 280)

    # Linhas de dados
    for r in rows_data:
        drow = tbl.rows[ri]; ri += 1
        cp(drow.cells[0], r.get('ref', ''), font='Calibri', align=CENTER)
        set_borders(drow.cells[0], top='nil', bottom='single', left='single', right='single')
        if htype == 'full':
            for ci, key in [(1, 'w'), (2, 'h'), (3, 'd'), (4, 't')]:
                cp(drow.cells[ci], r.get(key, ''), font='Calibri', align=CENTER)
                set_borders(drow.cells[ci], top='nil', bottom='single', left='nil', right='single')
        elif htype == 'nodiscount':
            for ci, key in [(1, 'w'), (2, 'h'), (3, 't')]:
                cp(drow.cells[ci], r.get(key, ''), font='Calibri', align=CENTER)
                set_borders(drow.cells[ci], top='nil', bottom='single', left='nil', right='single')
        else:
            cp(drow.cells[1], r.get('t', ''), font='Calibri', align=CENTER)
            set_borders(drow.cells[1], top='nil', bottom='single', left='nil', right='single')
        set_row_h(drow, 240)

    # Subtotal
    if has_sub:
        srow = tbl.rows[ri]; ri += 1
        if htype == 'full':
            sub_d = sum(float(r.get('d', '0').replace(',', '.')) for r in rows_data)
            sub_t = sum(float(r.get('t', '0').replace(',', '.')) for r in rows_data)
            srow.cells[0].merge(srow.cells[2])
            cp(srow.cells[0], 'Subtotal', bold=True, font='Calibri', align=CENTER)
            set_shd(srow.cells[0], GRAY_MID)
            set_borders(srow.cells[0], top='single', bottom='single',
                        left='single', right='single', right_color=BLACK)
            cp(srow.cells[3], fmt(sub_d), font='Calibri', align=CENTER)
            set_shd(srow.cells[3], WHITE)
            set_borders(srow.cells[3], top='nil', bottom='single', left='nil', right='single')
            cp(srow.cells[4], fmt(sub_t), font='Calibri', align=CENTER)
            set_borders(srow.cells[4], top='nil', bottom='single', left='nil', right='single')
        else:
            sub_t = sum(float(r.get('t', '0').replace(',', '.')) for r in rows_data)
            srow.cells[0].merge(srow.cells[2])
            cp(srow.cells[0], 'Subtotal', bold=True, font='Calibri', align=CENTER)
            set_shd(srow.cells[0], GRAY_MID)
            set_borders(srow.cells[0], top='single', bottom='single',
                        left='single', right='single', right_color=BLACK)
            cp(srow.cells[3], fmt(sub_t), font='Calibri', align=CENTER)
            set_borders(srow.cells[3], top='nil', bottom='single', left='nil', right='single')
        set_row_h(srow, 240)

    # Total
    trow2 = tbl.rows[ri]
    if htype == 'full':
        trow2.cells[0].merge(trow2.cells[3])
        cp(trow2.cells[0], final_lbl, bold=True, font='Calibri', align=CENTER)
        set_shd(trow2.cells[0], GRAY_MID)
        set_borders(trow2.cells[0], top='single', bottom='single',
                    left='single', right='single', right_color=BLACK)
        cp(trow2.cells[4], fmt(total_val), font='Calibri', align=CENTER)
        set_shd(trow2.cells[4], GRAY_MID)
        set_borders(trow2.cells[4], top='nil', bottom='single', left='nil', right='single')
    elif htype == 'nodiscount':
        trow2.cells[0].merge(trow2.cells[2])
        cp(trow2.cells[0], final_lbl, bold=True, font='Calibri', align=CENTER)
        set_shd(trow2.cells[0], GRAY_MID)
        set_borders(trow2.cells[0], top='single', bottom='single',
                    left='single', right='single', right_color=BLACK)
        cp(trow2.cells[3], fmt(total_val), font='Calibri', align=CENTER)
        set_shd(trow2.cells[3], GRAY_MID)
        set_borders(trow2.cells[3], top='nil', bottom='single', left='nil', right='single')
    else:
        cp(trow2.cells[0], final_lbl, bold=True, font='Calibri', align=CENTER)
        set_shd(trow2.cells[0], GRAY_MID)
        set_borders(trow2.cells[0], top='single', bottom='single',
                    left='single', right='single', right_color=BLACK)
        cp(trow2.cells[1], fmt(total_val), font='Calibri', align=CENTER)
        set_shd(trow2.cells[1], GRAY_MID)
        set_borders(trow2.cells[1], top='nil', bottom='single', left='nil', right='single')
    set_row_h(trow2, 240)

    set_widths(tbl, widths)


def add_itens_table(doc, code, total_val, items_info):
    desc_full = items_info.get(code, {}).get('desc', '')
    unit_str  = items_info.get(code, {}).get('unit', '')

    # Parágrafo separador entre calc e Itens (before=40, after=40 twips)
    gap_para(doc, before_twips=40, after_twips=40)

    itbl = doc.add_table(rows=2, cols=4)
    itbl.style = 'Table Grid'

    # Cabeçalho "Itens" — cinza escuro, Aptos Narrow, mesclado
    ih = itbl.rows[0]
    for ci in range(1, 4):
        ih.cells[0].merge(ih.cells[ci])
    cp(ih.cells[0], 'Itens', bold=True, color=BLACK, font='Aptos Narrow', align=CENTER)
    set_shd(ih.cells[0], GRAY_DARK)
    set_borders(ih.cells[0], top='single', bottom='single',
                left='single', right='single', right_color=BLACK)
    set_row_h(ih, 240)

    # Linha de dados: código | descrição | quantidade | unidade
    dr = itbl.rows[1]
    cp(dr.cells[0], code,          font='Calibri',      align=CENTER)
    cp(dr.cells[1], desc_full,     font='Aptos Narrow', align=CENTER)
    cp(dr.cells[2], fmt(total_val), font='Aptos Narrow', align=CENTER)
    cp(dr.cells[3], unit_str,      font='Aptos Narrow', align=CENTER)
    set_borders(dr.cells[0], top='nil', bottom='single', left='single', right='single')
    set_borders(dr.cells[1], top='nil', bottom='single', left='nil', right='single')
    set_borders(dr.cells[2], top='nil', bottom='single', left='nil', right='single')
    set_borders(dr.cells[3], top='nil', bottom='single', left='nil', right='single')
    set_row_h(dr, 240)

    set_widths(itbl, WIDTHS['itens'])


# ── Main ──────────────────────────────────────────────────────────────────────
def main(md_path, out_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    memorial_start = find_memorial_start(lines)
    body_lines = lines[:memorial_start] if memorial_start else lines

    occurrences = parse_calc_occurrences(body_lines)
    items_info, _ = parse_itens_tables(body_lines)
    consolidated = consolidate_occurrences(occurrences)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(21.0); sec.page_height  = Cm(29.7)
    sec.left_margin  = Cm(2.5);  sec.right_margin = Cm(2.5)
    sec.top_margin   = Cm(2.5);  sec.bottom_margin = Cm(2.5)

    # Título da seção
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run('- Memorial de cálculo e itens:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    for code, d in consolidated.items():
        add_calc_table(doc, d, items_info)
        add_itens_table(doc, code, d['total_sum'], items_info)
        # Parágrafo separador entre blocos (after=120 twips)
        gap_para(doc, before_twips=0, after_twips=120)

    doc.save(out_path)
    print(f'✅ Memorial salvo: {out_path}')
    print(f'   {len(consolidated)} itens únicos consolidados')
    print(f'⚠️  Execute: python3 fix_xml.py {out_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Uso: python3 gerar_memorial.py <input.md> <output.docx>')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2])
