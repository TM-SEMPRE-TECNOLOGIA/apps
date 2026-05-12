import sys
from docx import Document

def analyze_docx(path):
    try:
        doc = Document(path)
        print(f"Doc: {path}")
        print(f"Parágrafos: {len(doc.paragraphs)}")
        print(f"Tabelas: {len(doc.tables)}")
        
        print("\n--- Possíveis Ambientes (Cabeçalhos) ---")
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text and (any(x in text.upper() for x in ["DADOS DA DEPENDÊNCIA", "AMBIENTE:", "SETOR:"]) or text.isupper()):
                if len(text) < 100:
                    print(f"P{i}: {text}")

        print("\n--- Primeiras 5 Tabelas ---")
        for i, t in enumerate(doc.tables[:5]):
            rows = len(t.rows)
            cols = len(t.columns)
            print(f"Table {i}: {rows}x{cols}")
            for r_i, row in enumerate(t.rows[:3]):
                row_texts = [c.text.strip().replace('\n', ' ')[:50] for c in row.cells]
                print(f"   R{r_i}: {row_texts}")

        print("\n--- Últimas 5 Tabelas ---")
        for i, t in enumerate(doc.tables[-5:]):
            rows = len(t.rows)
            cols = len(t.columns)
            print(f"Final Table {i}: {rows}x{cols}")
            for r_i, row in enumerate(t.rows[:5]):
                row_texts = [c.text.strip().replace('\n', ' ')[:50] for c in row.cells]
                print(f"   R{r_i}: {row_texts}")

    except Exception as e:
        print(f"Erro ao ler docx: {e}")

if __name__ == "__main__":
    analyze_docx(sys.argv[1])
