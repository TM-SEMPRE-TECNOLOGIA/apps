import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_textbox(paragraph, text, width_cm=3.5, height_cm=0.6, left_cm=0, top_cm=0):
    """
    Injeta uma Caixa de Texto (Shape) real no parágrafo usando XML bruto.
    """
    run = paragraph.add_run()
    
    # Conversão de CM para EMUs (1 cm = 360000 EMUs)
    width_emu = int(width_cm * 360000)
    height_emu = int(height_cm * 360000)
    left_emu = int(left_cm * 360000)
    top_emu = int(top_cm * 360000)

    # XML para w:drawing com wp:anchor (flutuante)
    drawing_xml = f"""
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingdrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:posOffset>{left_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>{top_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="Text Box 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr/>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{width_emu}" cy="{height_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:solidFill>
                  <a:srgbClr val="C00000"/>
                </a:solidFill>
                <a:ln>
                    <a:noFill/>
                </a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:pPr>
                      <w:jc w:val="center"/>
                    </w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                        <w:b/>
                        <w:color w:val="FFFFFF"/>
                        <w:sz w:val="22"/>
                      </w:rPr>
                      <w:t>{text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr vert="horz" lIns="45720" tIns="22860" rIns="45720" bIns="22860" anchor="ctr">
                <a:spAutoFit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
    """
    
    drawing = parse_xml(drawing_xml)
    run._r.append(drawing)

if __name__ == "__main__":
    doc = Document()
    p = doc.add_paragraph("Teste de Imagem com Textbox Flutuante")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_img = p_img.add_run("[IMAGEM AQUI]")
    
    # Adicionando a Caixa de Texto
    create_textbox(p_img, "2,50 x 3,10 m", left_cm=10, top_cm=0)
    
    output = "test_textbox_xml.docx"
    doc.save(output)
    print(f"Documento salvo em {output}")
