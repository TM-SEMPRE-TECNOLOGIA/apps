"""Teste de TextBox com altura corrigida."""
import sys, os, zipfile, re
sys.path.insert(0, '.')

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from word_utils import criar_legenda_vermelha

saida = r'C:\Users\thiag\TM-MEUS-APPS\0 - NEXT APPS\TM Relatorio teste\output_test'
os.makedirs(saida, exist_ok=True)
output_file = os.path.join(saida, 'TESTE_TEXTBOX_CORRIGIDO.docx')

doc = Document()
doc.add_heading('Teste TextBox Altura Corrigida', 0)

p1 = doc.add_paragraph('[IMAGEM 1 - 2,50 x 3,00]')
p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
criar_legenda_vermelha(doc, ['2,50 m', '3,00 m'], p1)

p2 = doc.add_paragraph('[IMAGEM 2 - 4,00 x 2,80 - Desconto 1,50m²]')
p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
criar_legenda_vermelha(doc, ['4,00 m', '2,80 m', 'Desconto 1,50m²'], p2)

p4 = doc.add_paragraph('[IMAGEM 4 - Emassamento geral]')
p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
criar_legenda_vermelha(doc, ['Emassamento geral'], p4)

doc.save(output_file)
print(f"OK! Arquivo: {output_file}")
print(f"Tamanho: {os.path.getsize(output_file)} bytes")

from docx import Document as D2
d = D2(output_file)
print(f"python-docx reabriu OK! Parágrafos: {len(d.paragraphs)}")
print("Abra no Word para conferir a altura!")
