import os
import re
import copy
import random
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from word_utils import (
    ALTURA_PADRAO, 
    LARGURA_MAX_3_COL, 
    LARGURA_MAX_2_COL, 
    analisar_imagem, 
    otimizar_layout, 
    aplicar_estilo
)
from utils_sp import formatar_moeda_texto, parse_legendas_completo

def set_cell_background(cell, color="D9D9D9"):
    """Atribui cor de fundo à célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def _gerar_id_unico():
    """Gera um ID numérico grande e único para wp:docPr, evitando colisões."""
    return str(random.randint(100000000, 2147483647))

def _gerar_anchor_id():
    """Gera um anchorId hexadecimal de 8 caracteres para wp14:anchorId."""
    return ''.join(random.choices('0123456789ABCDEF', k=8))

def inserir_textbox_vermelha(paragrafo_ancora, texto, offset_left_emu=0, offset_top_emu=0, largura_emu=720000, altura_emu=228600):
    """
    Insere uma caixa de texto vermelha (Shape real) no documento usando OXML.
    Baseado na engenharia reversa do DOCX RESPLENDOR.
    
    Parâmetros:
    - paragrafo_ancora: parágrafo onde ancorar a TextBox
    - texto: texto da legenda
    - offset_left_emu: posição horizontal em EMUs (1 cm = 360000 EMU)
    - offset_top_emu: posição vertical em EMUs relativa ao parágrafo
    - largura_emu: largura da TextBox em EMUs
    - altura_emu: altura da TextBox em EMUs
    """
    doc_pr_id = _gerar_id_unico()
    anchor_id = _gerar_anchor_id()
    edit_id = _gerar_anchor_id()
    
    # Namespaces necessários
    nsmap = {
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'o': 'urn:schemas-microsoft-com:office:office',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    }
    
    # XML template da TextBox (baseado no RESPLENDOR)
    xml_str = f'''<mc:AlternateContent xmlns:mc="{nsmap['mc']}"
        xmlns:wp="{nsmap['wp']}" xmlns:wp14="{nsmap['wp14']}"
        xmlns:w="{nsmap['w']}" xmlns:wps="{nsmap['wps']}"
        xmlns:a="{nsmap['a']}" xmlns:v="{nsmap['v']}"
        xmlns:o="{nsmap['o']}" xmlns:w14="{nsmap['w14']}">
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                     simplePos="0" relativeHeight="251659264" behindDoc="0"
                     locked="0" layoutInCell="1" allowOverlap="1"
                     wp14:anchorId="{anchor_id}" wp14:editId="{edit_id}">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
              <wp:posOffset>{offset_left_emu}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
              <wp:posOffset>{offset_top_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{largura_emu}" cy="{altura_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{doc_pr_id}" name="Caixa de Texto {doc_pr_id}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{largura_emu}" cy="{altura_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                    <a:solidFill>
                      <a:srgbClr val="C00000"/>
                    </a:solidFill>
                    <a:ln w="6350">
                      <a:solidFill>
                        <a:srgbClr val="000000"/>
                      </a:solidFill>
                    </a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p>
                        <w:pPr>
                          <w:jc w:val="center"/>
                          <w:rPr>
                            <w:b/>
                            <w:bCs/>
                          </w:rPr>
                        </w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                            <w:b/>
                            <w:bCs/>
                            <w:color w:val="FFFFFF"/>
                            <w:sz w:val="22"/>
                            <w:szCs w:val="22"/>
                          </w:rPr>
                          <w:t xml:space="preserve">{texto}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                              horzOverflow="overflow" vert="horz" wrap="none"
                              lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                              numCol="1" spcCol="0" rtlCol="0" fromWordArt="0"
                              anchor="t" anchorCtr="0" forceAA="0" compatLnSpc="1">
                    <a:prstTxWarp prst="textNoShape">
                      <a:avLst/>
                    </a:prstTxWarp>
                    <a:noAutofit/>
                  </wps:bodyPr>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
            <wp14:sizeRelH relativeFrom="margin">
              <wp14:pctWidth>0</wp14:pctWidth>
            </wp14:sizeRelH>
            <wp14:sizeRelV relativeFrom="margin">
              <wp14:pctHeight>0</wp14:pctHeight>
            </wp14:sizeRelV>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
      <mc:Fallback>
        <w:pict>
          <v:shapetype id="_x0000_t202" coordsize="21600,21600" o:spt="202"
                       path="m,l,21600r21600,l21600,xe">
            <v:stroke joinstyle="miter"/>
            <v:path gradientshapeok="t" o:connecttype="rect"/>
          </v:shapetype>
          <v:shape id="CaixaTexto_{doc_pr_id}" o:spid="_x0000_s{doc_pr_id}"
                   type="#_x0000_t202"
                   style="position:absolute;margin-left:0;margin-top:0;width:56.7pt;height:18pt;z-index:251659264;visibility:visible"
                   fillcolor="#C00000" strokecolor="black" strokeweight=".5pt">
            <v:textbox>
              <w:txbxContent>
                <w:p>
                  <w:pPr>
                    <w:jc w:val="center"/>
                    <w:rPr>
                      <w:b/>
                      <w:bCs/>
                    </w:rPr>
                  </w:pPr>
                  <w:r>
                    <w:rPr>
                      <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                      <w:b/>
                      <w:bCs/>
                      <w:color w:val="FFFFFF"/>
                      <w:sz w:val="22"/>
                      <w:szCs w:val="22"/>
                    </w:rPr>
                    <w:t xml:space="preserve">{texto}</w:t>
                  </w:r>
                </w:p>
              </w:txbxContent>
            </v:textbox>
          </v:shape>
        </w:pict>
      </mc:Fallback>
    </mc:AlternateContent>'''
    
    # Parse o XML e insere como um run no parágrafo ancora
    textbox_element = parse_xml(xml_str)
    
    # Cria um run (w:r) para conter o AlternateContent
    run_element = OxmlElement('w:r')
    run_element.append(textbox_element)
    
    # Adiciona ao parágrafo ancora
    paragrafo_ancora._p.append(run_element)

def criar_legenda_vermelha(doc, legenda_list, paragrafo_ancora):
    """
    Cria legendas vermelhas reais (Shapes/TextBoxes) ancoradas ao parágrafo da imagem.
    Usa OXML baseado na engenharia reversa do DOCX RESPLENDOR.
    Cada legenda é uma caixa de texto vermelha com texto branco, bold, Arial 11pt.
    """
    if not legenda_list:
        return

    # Larguras em EMUs (1 cm = 360000 EMU, 1 pt = 12700 EMU)
    # Para textos curtos (ex: "2,00 m") usa ~2cm, para longos usa mais
    offset_left = 360000  # 1cm da margem esquerda (posição inicial)
    
    for i, texto in enumerate(legenda_list):
        # Calcula largura baseada no comprimento do texto
        char_count = len(texto)
        if char_count <= 8:
            largura = 720000   # ~2cm
        elif char_count <= 15:
            largura = 1080000  # ~3cm
        else:
            largura = 1800000  # ~5cm
        
        altura = 259345  # ~0.68cm (altura real do RESPLENDOR)
        
        # Posição vertical: logo abaixo da imagem
        offset_top = -10000  # Levemente acima da próxima linha
        
        inserir_textbox_vermelha(
            paragrafo_ancora, 
            texto, 
            offset_left_emu=offset_left, 
            offset_top_emu=offset_top,
            largura_emu=largura,
            altura_emu=altura
        )
        
        # Avança a posição horizontal para a próxima caixa
        offset_left += largura + 72000  # Gap de ~0.2cm entre caixas

def inserir_conteudo_sp(modelo_path, conteudo, output_path, selected_description=None):
    doc = Document(modelo_path)

    # Substituição do placeholder {Desc_here} no documento inteiro (parágrafos e tabelas)
    if selected_description:
        # Substituir nos parágrafos
        for p in doc.paragraphs:
            if "{Desc_here}" in p.text:
                p.text = p.text.replace("{Desc_here}", selected_description)
        
        # Substituir nas tabelas (caso o placeholder esteja dentro de uma célula)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{Desc_here}" in p.text:
                            p.text = p.text.replace("{Desc_here}", selected_description)

    contador_imagens = 0
    paragrafo_insercao_index = None

    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("[AVISO] Marca '{{start_here}}' não encontrada.")
        return 0
    
    # Aplica otimização de imagens (reutiliza a lógica do Tradicional)
    conteudo_otimizado = otimizar_layout(conteudo)

    # Inverte ordem para inserir "antes" na posição
    conteudo_invertido = list(reversed(conteudo_otimizado))

    for item in conteudo_invertido:
        # Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip()
            nivel = item.count("»")
            
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            if nivel == 0:
                p.style = 'Heading 1'
                p.add_run(f"- {titulo}:")
            elif nivel == 1:
                p.style = 'Heading 2'
                p.add_run(f"{titulo}:")
            elif nivel == 2:
                p.style = 'Heading 3'
                p.add_run(f"{titulo}:")
            
        # Textos Padrão (- Detalhes:, - Vista ampla:)
        elif isinstance(item, dict) and "texto_padrao" in item:
            texto = f"- {item['texto_padrao']}:"
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(texto)
            aplicar_estilo(run, 11, True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Descrição Técnica Textual
        elif isinstance(item, dict) and "texto_descricao" in item:
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            texto_raw = item["texto_descricao"]
            
            # Processa o marcador <RED>
            partes = re.split(r'(<RED>.*?</RED>)', texto_raw)
            for parte in partes:
                if parte.startswith("<RED>") and parte.endswith("</RED>"):
                    conteudo_red = parte[5:-6]
                    run = p.add_run(conteudo_red)
                    run.font.color.rgb = RGBColor(255, 0, 0) # VERMELHO
                    aplicar_estilo(run, 11, True) # Negrito
                else:
                    run = p.add_run(parte)
                    aplicar_estilo(run, 11, False)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Tabela de Medição Otimizada
        elif isinstance(item, dict) and "tabela_medicao" in item:
            # Pular uma linha antes
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            tabela_dados = item["tabela_medicao"]
            
            if tabela_dados["tipo"] == "pintura":
                tabela = doc.add_table(rows=2, cols=5)
                doc.paragraphs[paragrafo_insercao_index]._p.addnext(tabela._tbl)
                
                tabela.style = 'Table Grid'
                tabela.autofit = False
                tabela.allow_autofit = False

                larguras = [3.5, 3.0, 3.0, 3.0, 3.5]
                for i, col in enumerate(tabela.columns):
                    col.width = Cm(larguras[i])

                # Linha de Título Geral (Mesclada)
                titulo_cel = tabela.cell(0, 0)
                for i in range(1, 4):
                    titulo_cel = titulo_cel.merge(tabela.cell(0, i))
                titulo_cel = titulo_cel.merge(tabela.cell(0, 4))
                
                titulo_cel.text = "Pintura em látex acrílica standard fosca sem emassamento, 3 demãos, com aplicação de selador para exterior (ITEM 17.4)"
                titulo_cel.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run_t = titulo_cel.paragraphs[0].runs[0]
                run_t.font.bold = True
                run_t.font.size = Pt(10)
                set_cell_background(titulo_cel, "808080")

                # Linha de Cabeçalho
                hdr_cells = tabela.rows[1].cells
                headers = ["REFERÊNCIA", "LARGURA (m)", "ALTURA (m)", "DESCONTO", "TOTAL (m²)"]
                for i, text in enumerate(headers):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)
                    set_cell_background(hdr_cells[i], "D9D9D9")

                soma_bruta = 0.0
                soma_desconto = 0.0

                for med in tabela_dados["medidas"]:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = med["referencia"]
                    row_cells[1].text = formatar_moeda_texto(med["largura"])
                    row_cells[2].text = formatar_moeda_texto(med["altura"])
                    row_cells[3].text = formatar_moeda_texto(med["desconto"])
                    row_cells[4].text = formatar_moeda_texto(med["subtotal"])
                    
                    for c in row_cells:
                        c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    soma_bruta += med["subtotal"]
                    soma_desconto += med["desconto"]

                # Linha Subtotal
                row_sub = tabela.add_row().cells
                sub_label = row_sub[0].merge(row_sub[1]).merge(row_sub[2])
                sub_label.text = "Subtotal"
                sub_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                sub_label.paragraphs[0].runs[0].font.bold = True
                set_cell_background(sub_label, "D9D9D9")
                
                row_sub[3].text = formatar_moeda_texto(soma_desconto)
                row_sub[4].text = formatar_moeda_texto(soma_bruta)
                for i in [3, 4]:
                    row_sub[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    set_cell_background(row_sub[i], "D9D9D9")

                # Linha Total
                row_total = tabela.add_row().cells
                total_label = row_total[0].merge(row_total[1]).merge(row_total[2]).merge(row_total[3])
                total_label.text = "Total"
                total_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                total_label.paragraphs[0].runs[0].font.bold = True
                set_cell_background(total_label, "D9D9D9")

                area_liquida = soma_bruta - soma_desconto
                row_total[4].text = formatar_moeda_texto(area_liquida)
                row_total[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row_total[4].paragraphs[0].runs[0].font.bold = True
                set_cell_background(row_total[4], "D9D9D9")

            elif tabela_dados["tipo"] == "mobiliario":
                tabela = doc.add_table(rows=2, cols=2)
                doc.paragraphs[paragrafo_insercao_index]._p.addnext(tabela._tbl)
                tabela.style = 'Table Grid'
                tabela.autofit = False
                tabela.width = Cm(16.0)
                
                # Titulo Mobiliario
                tit_mob = tabela.cell(0, 0).merge(tabela.cell(0, 1))
                tit_mob.text = "Deslocamento ou remanejamento de mobiliário dentro da agência (ITEM 13.12)"
                tit_mob.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tit_mob.paragraphs[0].runs[0].font.bold = True
                set_cell_background(tit_mob, "808080")

                # Header
                h_cells = tabela.rows[1].cells
                h_cells[0].text = "REFERÊNCIA"
                h_cells[1].text = "TOTAL (UN)"
                for c in h_cells: 
                    c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    c.paragraphs[0].runs[0].font.bold = True
                    set_cell_background(c, "D9D9D9")

                soma_un = 0.0
                for med in tabela_dados["medidas"]:
                    r_cells = tabela.add_row().cells
                    r_cells[0].text = med["referencia"]
                    r_cells[1].text = formatar_moeda_texto(med["total_un"])
                    for c in r_cells: c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    soma_un += med["total_un"]

                r_tot = tabela.add_row().cells
                r_tot[0].text = "Total"
                r_tot[1].text = formatar_moeda_texto(soma_un)
                for c in r_tot: 
                    c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    c.paragraphs[0].runs[0].font.bold = True
                    set_cell_background(c, "D9D9D9")

            # Espaçamento extra pós tabela
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')

        # Imagens Simples
        elif isinstance(item, dict) and "imagem" in item:
            from PIL import Image
            imagem_path = item["imagem"]
            if os.path.exists(imagem_path):
                try:
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                    run = p.add_run()
                    with Image.open(imagem_path) as img:
                        w, h = img.size
                        prop = ALTURA_PADRAO / (h / 37.7952755906)
                        w_cm = (w * prop) / 37.7952755906
                    run.add_picture(imagem_path, width=Cm(w_cm), height=Cm(ALTURA_PADRAO))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    contador_imagens += 1
                    
                    # Cria a legenda vermelha ancorada abaixo da imagem
                    legendas = parse_legendas_completo(os.path.basename(imagem_path))
                    criar_legenda_vermelha(doc, legendas, p)
                except Exception as e:
                    print(f"Erro imagem {imagem_path}: {e}")
                    
        # Tabela Imagens (Reaproveitado do util base)
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            tabela = doc.add_table(rows=1, cols=colunas)
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)
            
            w_col = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(w_col)
                
            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].add_run()
                    try:
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                        
                        # Legendas para imagens em grid
                        legendas = parse_legendas_completo(os.path.basename(img_path))
                        # Em tabelas, vamos tentar um offset menor para caber na coluna
                        # Ou um posicionamento diferente. Por enquanto, vamos usar o padrão.
                        criar_legenda_vermelha(doc, legendas, cell.paragraphs[0])
                    except Exception as e:
                        print(f"Erro ao inserir na tabela: {e}")
        
        # Quebra de Página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

    # Remover a marcação final para não sujar o relatório
    doc.paragraphs[paragrafo_insercao_index].text = doc.paragraphs[paragrafo_insercao_index].text.replace("{{start_here}}", "")
    
    doc.save(output_path)
    return contador_imagens
