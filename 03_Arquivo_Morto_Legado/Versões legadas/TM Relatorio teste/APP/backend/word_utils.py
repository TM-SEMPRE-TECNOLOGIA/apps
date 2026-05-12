import os
import datetime
import random
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image, UnidentifiedImageError

PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

# Constantes de Layout
ALTURA_PADRAO = 10.0  # cm
LARGURA_MAX_3_COL = 5.63  # cm
LARGURA_MAX_2_COL = 7.50  # cm

from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import RGBColor
from utils_sp import parse_legendas_completo

def set_cell_background(cell, color="D9D9D9"):
    """Atribui cor de fundo à célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def _gerar_id_unico():
    """Gera um ID numérico grande e único para wp:docPr."""
    return str(random.randint(100000000, 2147483647))

def _gerar_anchor_id():
    """Gera um anchorId hexadecimal de 8 caracteres."""
    return ''.join(random.choices('0123456789ABCDEF', k=8))

def inserir_textbox_vermelha(paragrafo_ancora, texto, offset_left_emu=0, offset_top_emu=0, largura_emu=720000, altura_emu=228600):
    """
    Insere uma caixa de texto vermelha (Shape real) no documento usando OXML.
    Baseado na engenharia reversa do DOCX RESPLENDOR.
    """
    doc_pr_id = _gerar_id_unico()
    anchor_id = _gerar_anchor_id()
    edit_id = _gerar_anchor_id()
    
    nsmap = {
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'o': 'urn:schemas-microsoft-com:office:office',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    }
    
    xml_str = f'''<mc:AlternateContent xmlns:mc="{nsmap['mc']}"
        xmlns:wp="{nsmap['wp']}" xmlns:wp14="{nsmap['wp14']}"
        xmlns:w="{nsmap['w']}" xmlns:wps="{nsmap['wps']}"
        xmlns:a="{nsmap['a']}" xmlns:v="{nsmap['v']}"
        xmlns:o="{nsmap['o']}" xmlns:w14="{nsmap['w14']}">
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                     simplePos="0" relativeHeight="251659264" behindDoc="0"
                     locked="0" layoutInCell="1" allowOverlap="1"
                     wp14:anchorId="{anchor_id}" wp14:editId="{edit_id}">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
              <wp:posOffset>{offset_left_emu}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
              <wp:posOffset>{offset_top_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{largura_emu}" cy="{altura_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{doc_pr_id}" name="Caixa de Texto {doc_pr_id}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{largura_emu}" cy="{altura_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                    <a:solidFill>
                      <a:srgbClr val="C00000"/>
                    </a:solidFill>
                    <a:ln w="6350">
                      <a:solidFill>
                        <a:srgbClr val="000000"/>
                      </a:solidFill>
                    </a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p>
                        <w:pPr>
                          <w:jc w:val="center"/>
                          <w:rPr>
                            <w:b/>
                            <w:bCs/>
                          </w:rPr>
                        </w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                            <w:b/>
                            <w:bCs/>
                            <w:color w:val="FFFFFF"/>
                            <w:sz w:val="22"/>
                            <w:szCs w:val="22"/>
                          </w:rPr>
                          <w:t xml:space="preserve">{texto}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                              horzOverflow="overflow" vert="horz" wrap="none"
                              lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                              numCol="1" spcCol="0" rtlCol="0" fromWordArt="0"
                              anchor="t" anchorCtr="0" forceAA="0" compatLnSpc="1">
                    <a:prstTxWarp prst="textNoShape">
                      <a:avLst/>
                    </a:prstTxWarp>
                    <a:noAutofit/>
                  </wps:bodyPr>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
            <wp14:sizeRelH relativeFrom="margin">
              <wp14:pctWidth>0</wp14:pctWidth>
            </wp14:sizeRelH>
            <wp14:sizeRelV relativeFrom="margin">
              <wp14:pctHeight>0</wp14:pctHeight>
            </wp14:sizeRelV>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
      <mc:Fallback>
        <w:pict>
          <v:shapetype id="_x0000_t202" coordsize="21600,21600" o:spt="202"
                       path="m,l,21600r21600,l21600,xe">
            <v:stroke joinstyle="miter"/>
            <v:path gradientshapeok="t" o:connecttype="rect"/>
          </v:shapetype>
          <v:shape id="CaixaTexto_{doc_pr_id}" o:spid="_x0000_s{doc_pr_id}"
                   type="#_x0000_t202"
                   style="position:absolute;margin-left:0;margin-top:0;width:56.7pt;height:18pt;z-index:251659264;visibility:visible"
                   fillcolor="#C00000" strokecolor="black" strokeweight=".5pt">
            <v:textbox>
              <w:txbxContent>
                <w:p>
                  <w:pPr>
                    <w:jc w:val="center"/>
                    <w:rPr>
                      <w:b/>
                      <w:bCs/>
                    </w:rPr>
                  </w:pPr>
                  <w:r>
                    <w:rPr>
                      <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                      <w:b/>
                      <w:bCs/>
                      <w:color w:val="FFFFFF"/>
                      <w:sz w:val="22"/>
                      <w:szCs w:val="22"/>
                    </w:rPr>
                    <w:t xml:space="preserve">{texto}</w:t>
                  </w:r>
                </w:p>
              </w:txbxContent>
            </v:textbox>
          </v:shape>
        </w:pict>
      </mc:Fallback>
    </mc:AlternateContent>'''
    
    textbox_element = parse_xml(xml_str)
    run_element = OxmlElement('w:r')
    run_element.append(textbox_element)
    paragrafo_ancora._p.append(run_element)

def criar_legenda_vermelha(doc, legenda_list, paragrafo_ancora):
    """
    Cria legendas vermelhas reais (Shapes/TextBoxes) ancoradas ao parágrafo da imagem.
    Usa OXML baseado na engenharia reversa do DOCX RESPLENDOR.
    """
    if not legenda_list:
        return

    offset_left = 360000  # 1cm da margem esquerda
    
    for i, texto in enumerate(legenda_list):
        char_count = len(texto)
        if char_count <= 8:
            largura = 720000   # ~2cm
        elif char_count <= 15:
            largura = 1080000  # ~3cm
        else:
            largura = 1800000  # ~5cm
        
        altura = 259345  # ~0.68cm (altura real do RESPLENDOR)
        offset_top = -10000
        
        inserir_textbox_vermelha(
            paragrafo_ancora, 
            texto, 
            offset_left_emu=offset_left, 
            offset_top_emu=offset_top,
            largura_emu=largura,
            altura_emu=altura
        )
        
        offset_left += largura + 72000  # Gap de ~0.2cm entre caixas




def cm_to_px(cm):
    return cm * 37.7952755906


def analisar_imagem(caminho_imagem):
    """
    Analisa a imagem e retorna suas características para layout.
    Retorna: (is_vertical, largura_em_10cm)
    """
    try:
        with Image.open(caminho_imagem) as img:
            w, h = img.size
            aspect_ratio = w / h
            largura_em_10cm = ALTURA_PADRAO * aspect_ratio
            
            # Definição de Vertical: Altura > Largura
            # E também não pode ser absurdamente larga (ex: panorâmica vertical)
            is_vertical = h > w
            
            return is_vertical, largura_em_10cm
    except Exception:
        return False, 999.0  # Trata como horizontal/erro em caso de falha


def otimizar_layout(conteudo_original):
    """
    Reorganiza o conteúdo agrupando imagens verticais em tabelas.
    Estratégia Greedy: Tenta agrupar 3, se não der tenta 2, senão 1.
    """
    novo_conteudo = []
    buffer_imagens = []

    def processar_buffer():
        nonlocal buffer_imagens
        while buffer_imagens:
            qtd = len(buffer_imagens)
            grupo_formado = False
            
            # Tenta formar grupo de 3
            if qtd >= 3:
                candidatos = buffer_imagens[:3]
                # Verifica se TODOS cabem na largura de 3 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_3_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 3
                    })
                    buffer_imagens = buffer_imagens[3:]
                    grupo_formado = True
            
            # Se não formou de 3, tenta de 2
            if not grupo_formado and qtd >= 2:
                candidatos = buffer_imagens[:2]
                # Verifica se TODOS cabem na largura de 2 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_2_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 2
                    })
                    buffer_imagens = buffer_imagens[2:]
                    grupo_formado = True
            
            # Se não formou grupo (ou sobrou 1), processa individualmente
            if not grupo_formado:
                img = buffer_imagens.pop(0)
                # Se for "Vertical" mas não coube nos grupos, vira Horizontal (Single)
                # Mantém como imagem normal para ser inserida centralizada
                novo_conteudo.append({"imagem": img['caminho']})

    for item in conteudo_original:
        if isinstance(item, dict) and "imagem" in item:
            caminho = item["imagem"]
            is_vert, larg_10cm = analisar_imagem(caminho)
            
            if is_vert:
                buffer_imagens.append({
                    "caminho": caminho,
                    "largura_10cm": larg_10cm
                })
            else:
                # Se veio uma horizontal, processa o que estava acumulado antes
                processar_buffer()
                novo_conteudo.append(item)
        else:
            # Títulos, quebras, etc. Processa buffer antes.
            processar_buffer()
            novo_conteudo.append(item)

    # Processa qualquer sobra final
    processar_buffer()
    
    return novo_conteudo



def aplicar_estilo(run, size_pt, is_bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = is_bold


def inserir_conteudo(modelo_path, conteudo, output_path, selected_description=None):
    doc = Document(modelo_path)

    # Substituição do placeholder {Desc_here} no documento inteiro (parágrafos e tabelas)
    if selected_description:
        # Substituir nos parágrafos
        for p in doc.paragraphs:
            if "{Desc_here}" in p.text:
                p.text = p.text.replace("{Desc_here}", selected_description)
        
        # Substituir nas tabelas (caso o placeholder esteja dentro de uma célula)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{Desc_here}" in p.text:
                            p.text = p.text.replace("{Desc_here}", selected_description)

    contador_imagens = 0
    paragrafo_insercao_index = None

    # Localiza a marca de inserção
    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("[AVISO] Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    # Otimiza o layout (agrupamento de imagens)
    conteudo_otimizado = otimizar_layout(conteudo)

    # Processa o conteúdo (reverso para manter ordem de inserção no topo)
    conteudo_invertido = list(reversed(conteudo_otimizado))

    for item in conteudo_invertido:
        # 🟦 Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, True)

        # 🟩 Imagens (Individuais / Horizontais)
        elif isinstance(item, dict) and "imagem" in item:
            imagem_path = item["imagem"]

            if os.path.exists(imagem_path) and os.path.getsize(imagem_path) > 0:
                try:
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size

                        # Altura fixa em 10 cm e largura proporcional
                        altura_desejada_cm = ALTURA_PADRAO
                        proporcao = altura_desejada_cm / (altura_original / 37.7952755906)
                        largura_proporcional_cm = largura_original * proporcao / 37.7952755906

                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        run = p.add_run()
                        run.add_picture(
                            imagem_path,
                            width=Cm(largura_proporcional_cm),
                            height=Cm(altura_desejada_cm)
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1
                        
                        # Cria a legenda vermelha ancorada abaixo da imagem
                        legendas = parse_legendas_completo(os.path.basename(imagem_path))
                        criar_legenda_vermelha(doc, legendas, p)

                except UnidentifiedImageError:
                    msg = f"[ERRO: formato não reconhecido] {imagem_path}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                except Exception as e:
                    msg = f"[ERRO ao inserir imagem: {imagem_path}] {e}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            else:
                msg = f"[ERRO: imagem inválida] {imagem_path}"
                print(msg)
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 🟦 Tabela de Imagens (Agrupadas)
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            
            # Insere parágrafo antes da tabela para separação
            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            # Cria tabela
            tabela = doc.add_table(rows=1, cols=colunas)
            
            # Move a tabela para depois do parágrafo de separação
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            tabela.allow_autofit = False
            
            # Remove bordas (opcional, mas recomendado para layout limpo)
            # Para remover bordas, iteramos sobre as bordas ou usamos estilo padrão sem borda
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)

            # Define largura das colunas (distribuição igual)
            # Largura total página A4 (21cm) - Margens (3cm esq + 2cm dir = 5cm) = 16cm útil
            # Ajuste fino: 16cm / colunas
            largura_coluna = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(largura_coluna)

            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    
                    try:
                        # Insere com altura fixa de 10cm
                        # A largura será proporcional. Como já validamos no otimizar_layout,
                        # sabemos que ela cabe na coluna!
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                        
                        # Legendas para imagens em grid
                        legendas = parse_legendas_completo(os.path.basename(img_path))
                        criar_legenda_vermelha(doc, legendas, cell.paragraphs[0])
                    except Exception as e:
                        paragraph.add_run(f"[Erro img: {os.path.basename(img_path)}]")
                        print(f"Erro ao inserir na tabela: {e}")

        # 🟨 Quebra de página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

        # 🟩 Novo tipo: Parágrafo
        elif isinstance(item, dict) and "paragrafo" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')

    # Salva documento final
    doc.save(output_path)
    print(f"[OK] Imagens redimensionadas para 10 cm de altura ({contador_imagens} inseridas).")
    return contador_imagens
