import os
import shutil
import base64
from io import BytesIO
from typing import List, Optional, Dict
from fastapi import FastAPI, Request, File, UploadFile, HTTPException, Query
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
from PIL import Image
import tkinter as tk
from tkinter import filedialog

# Reaproveitar lógica de geração atual
from generator import build_content_from_root, generate_report
from llm_generator import call_llm_with_services

# SP2 — Modo São Paulo 2 (Contrato 1565 — São José do Rio Preto)
from generator_sp2 import build_content_sp2
from word_utils_sp2 import inserir_conteudo_sp2

# [NEW] Sistema de placeholders dinâmicos (AutoRelatorio v3.2)
from routes import router as placeholders_router

app = FastAPI(title="TM Relatório API")

# Configuração de CORS para permitir requisições do Next.js (normalmente porta 3000) e Vite (5173)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Em produção, restringir para os domínios necessários
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# [NEW] Incluir router de placeholders dinâmicos (AutoRelatorio v3.2)
app.include_router(placeholders_router)

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(PROJECT_DIR, "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

class ScanRequest(BaseModel):
    pasta_raiz: str
    pasta_saida: Optional[str] = None
    tipo_relatorio: str = "tradicional"
    reading_mode: str = "disco"  # "disco" (pastas estruturadas) ou "app" (pasta plana)

class GenerateRequest(BaseModel):
    pasta_raiz: str
    modelo: str
    pasta_saida: str
    conteudo: list
    tipo_relatorio: str = "tradicional"
    selected_description_key: Optional[str] = None
    modal_data: Optional[Dict[str, list]] = None
    annotated_images: Optional[Dict[str, str]] = None
    # SP2: dados do cabeçalho institucional (contrato, OS, agência, etc.)
    meta_sp2: Optional[Dict[str, str]] = None
    # Metadados dinâmicos para todos os modos (substituição de placeholders {campo} no template)
    meta_fields: Optional[Dict[str, str]] = None

# Descrições dinâmicas para o placeholder {Desc_here}
DESCRICOES_OPCOES = {
    "Desc 1": "Descrição do serviço: Informamos que foi realizada visita técnica à agência para fins de levantamento preventivo. A atividade contou com o acompanhamento do gerente [Nome do Gerente] – Matrícula [Número da Matrícula], que esteve presente durante todo o procedimento. Durante a vistoria, foram identificadas necessidades de intervenção nos seguintes itens.",
    "Desc 2": "Descrição do serviço: No cumprimento das atividades programadas, nosso técnico realizou uma visita à agência para a execução do levantamento preventivo. O gerente [Nome do Gerente] – Matrícula [Número da Matrícula] acompanhou todas as etapas do procedimento. Após a avaliação, verificou-se a necessidade de intervenção nos seguintes itens.",
    "Desc 3": "Descrição do serviço: Informamos que nosso técnico realizou uma visita à agência para a execução do levantamento preventivo. Durante a visita, o gerente [Nome do Gerente] – Matrícula [Número da Matrícula] esteve presente, acompanhando todo o procedimento. Constatou-se a necessidade de intervenção nos seguintes itens.",
    "Desc 4": "Descrição do serviço: Nosso técnico realizou uma visita à agência para a execução do levantamento preventivo. O gerente [Nome do Gerente] – Matrícula [Número da Matrícula] acompanhou todo o processo. Durante a vistoria, identificamos a necessidade de intervenção nos seguintes itens."
}

class OpenFolderRequest(BaseModel):
    path: str

@app.get("/api/templates")
async def list_templates():
    try:
        files = [f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx")]
        files.sort(key=str.casefold)
        return {"templates": files}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/template-placeholders")
async def get_template_placeholders(template: str = Query(...)):
    """Lê um template DOCX e retorna todos os placeholders {campo} encontrados,
    excluindo os que são controlados internamente (start_here, Desc_here)."""
    import re
    from docx import Document as DocxDocument

    INTERNAL = {"start_here", "Desc_here", "desc_here"}
    template_path = os.path.join(TEMPLATES_DIR, template)
    if not os.path.isfile(template_path):
        raise HTTPException(status_code=404, detail="Template não encontrado")

    try:
        doc = DocxDocument(template_path)
        found: list[str] = []
        seen: set[str] = set()
        pattern = re.compile(r"\{(\w+)\}")

        def scan_text(text: str):
            for m in pattern.finditer(text):
                key = m.group(1)
                if key not in INTERNAL and key not in seen:
                    seen.add(key)
                    found.append(key)

        # Parágrafos do corpo
        for para in doc.paragraphs:
            scan_text(para.text)

        # Tabelas do corpo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_text(cell.text)

        # Cabeçalho e rodapé de todas as seções
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                scan_text(hdr_para.text)
            for hdr_table in section.header.tables:
                for row in hdr_table.rows:
                    for cell in row.cells:
                        scan_text(cell.text)
            for ftr_para in section.footer.paragraphs:
                scan_text(ftr_para.text)

        return {"placeholders": found}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/upload-template")
async def upload_template(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Apenas arquivos .docx são permitidos")
    
    try:
        filepath = os.path.join(TEMPLATES_DIR, file.filename)
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return {"message": "Template salvo", "filename": file.filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/scan")
async def scan_directory(data: ScanRequest):
    pasta_raiz = data.pasta_raiz.strip()
    
    if not pasta_raiz or not os.path.isdir(pasta_raiz):
        raise HTTPException(status_code=400, detail="Pasta raiz inválida")
    
    pasta_saida = data.pasta_saida or os.path.join(PROJECT_DIR, "..", "..", "DOCUMENTOS", "output")
    os.makedirs(pasta_saida, exist_ok=True)
    log_errors = os.path.join(pasta_saida, "erros_pastas.txt")
    
    def dummy_logger(msg): pass
    
    try:
        # reading_mode define HOW to read folders (disco=estruturado ou app=plano)
        # tipo_relatorio define HOW to generate Word (sp, sp2, tradicional)

        if data.reading_mode == "app":
            # Modo APP: lê todas as imagens de forma plana
            from generator_app import build_content_app
            conteudo = build_content_app(pasta_raiz, log_errors, logger=dummy_logger)
        elif data.tipo_relatorio == "sp":
            # Modo DISCO + SP: lê pastas com estrutura SP
            from generator_sp import build_content_sp
            conteudo = build_content_sp(pasta_raiz, log_errors, logger=dummy_logger)
        elif data.tipo_relatorio == "sp2":
            # Modo DISCO + SP2: lê pastas com estrutura SP2
            conteudo = build_content_sp2(pasta_raiz, log_errors, logger=dummy_logger)
        else:
            # Modo DISCO + Tradicional: lê pastas com estrutura tradicional
            conteudo = build_content_from_root(pasta_raiz, log_errors, logger=dummy_logger)
        return {"conteudo": conteudo}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/image")
async def get_image_full(path: str = Query(...)):
    """Serve a imagem original em resolução completa para o editor modal."""
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    # Serve o arquivo diretamente — sem redimensionar, sem comprimir
    return FileResponse(path, media_type="image/jpeg", headers={
        "Cache-Control": "no-store",
        "X-Image-Source": "original"
    })

@app.get("/api/thumbnail")
async def get_thumbnail(path: str = Query(...)):
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
        
    try:
        img = Image.open(path)
        img.thumbnail((200, 200), Image.Resampling.LANCZOS)
        
        buffered = BytesIO()
        img_format = img.format if img.format else 'JPEG'
        
        if img.mode in ('RGBA', 'P') and img_format in ('JPEG', 'JPG'):
             img = img.convert('RGB')
             
        img.save(buffered, format=img_format)
        buffered.seek(0)
        
        from starlette.responses import StreamingResponse
        return StreamingResponse(buffered, media_type=f"image/{img_format.lower()}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate")
async def generate_docx(data: GenerateRequest):
    pasta_raiz = data.pasta_raiz.strip()
    modelo_nome = data.modelo.strip()
    pasta_saida = data.pasta_saida.strip()
    conteudo_editado = data.conteudo
    
    # Resolve a descrição baseada na chave enviada pelo frontend
    selected_description = None
    if data.selected_description_key:
        selected_description = DESCRICOES_OPCOES.get(data.selected_description_key)
    
    if not all([pasta_raiz, modelo_nome, pasta_saida, conteudo_editado]):
        raise HTTPException(status_code=400, detail="Faltam parâmetros")
        
    modelo_path = os.path.join(TEMPLATES_DIR, modelo_nome)
    if not os.path.isfile(modelo_path):
        raise HTTPException(status_code=404, detail=f"Modelo não encontrado: {modelo_nome}")
        
    try:
        pasta_saida = os.path.abspath(pasta_saida)
        os.makedirs(pasta_saida, exist_ok=True)
        nome_pasta_raiz = os.path.basename(pasta_raiz.strip(os.sep))
        output_docx = os.path.join(pasta_saida, f"RELATÓRIO FOTOGRÁFICO - {nome_pasta_raiz} - LEVANTAMENTO PREVENTIVO.docx")
        
        log_errors = os.path.join(pasta_saida, "erros_pastas.txt")
        log_process = os.path.join(pasta_saida, "process_log.txt")
        
        def file_logger(msg: str) -> None:
            with open(log_process, "a", encoding="utf-8") as f:
                f.write(msg + "\n")

        with open(log_process, "w", encoding="utf-8") as f:
            f.write("PROCESS LOG - Geração de Relatório via Web (FastAPI)\n\n")
            
        file_logger(f"Pasta raiz: {pasta_raiz}")
        file_logger(f"Modelo: {modelo_path}")
        file_logger(f"Tipo: {data.tipo_relatorio}")
        file_logger(f"Saída: {pasta_saida}\n")
        
        # Processamento de modal_data e injeção de tabelas
        conteudo_final = []
        modal_data = data.modal_data or {}

        # ── Decodificar imagens anotadas (base64 → temp PNG) ──
        import tempfile
        temp_files: list[str] = []
        path_map: dict[str, str] = {}

        if data.annotated_images:
            for orig_path, data_url in data.annotated_images.items():
                try:
                    # Remove o header "data:image/png;base64," da URL
                    if ',' in data_url:
                        _, encoded = data_url.split(',', 1)
                    else:
                        encoded = data_url

                    decoded = base64.b64decode(encoded)
                    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    tmp.write(decoded)
                    tmp.close()
                    temp_files.append(tmp.name)
                    path_map[orig_path] = tmp.name
                    file_logger(f"[OK] Imagem anotada decodificada: {orig_path} → {tmp.name}")
                except Exception as e:
                    file_logger(f"[AVISO] Falha ao decodificar imagem anotada {orig_path}: {e}")

        # Substituir caminhos das imagens anotadas no conteúdo
        conteudo_editado = [
            {**item, 'imagem': path_map.get(item['imagem'], item['imagem'])}
            if isinstance(item, dict) and 'imagem' in item else item
            for item in conteudo_editado
        ]

        # Variáveis para agregação no modo Tradicional
        # A chave será o id do item, e o valor dict(item_obj, medicoes[], total)
        agregador_tradicional = {}
        desc_section = ''  # observação técnica da seção atual
        
        def flush_tradicional():
            nonlocal desc_section
            if desc_section:
                conteudo_final.append({"paragrafo": True})
                conteudo_final.append({"descricao_texto": desc_section})
                desc_section = ''
            for item_id, dados in agregador_tradicional.items():
                is_area = dados['un'] in ['m²', 'm']

                conteudo_final.append({"paragrafo": True})
                conteudo_final.append({
                    "tabela_calculo": { "medicoes": dados["medicoes"], "is_area": is_area }
                })
                conteudo_final.append({"paragrafo": True})
                conteudo_final.append({
                    "tabela_itens": {
                        "id": item_id,
                        "desc": dados["desc"],
                        "total": dados["total_acumulado"],
                        "un": dados["un"],
                        "totalLabel": "Total"
                    }
                })
            agregador_tradicional.clear()

        for item in conteudo_editado:
            # Se encontrar uma string no modo tradicional que seja MAIN section (sem "»"), fazer flush
            # Detalhes subsections têm "»" e não devem acionar flush
            if data.tipo_relatorio == "tradicional" and isinstance(item, str) and "»" not in item:
                flush_tradicional()

            conteudo_final.append(item)
            
            if isinstance(item, dict) and 'imagem' in item:
                img_path = item['imagem']
                services = modal_data.get(img_path)
                is_fachada = item.get('imagem_fachada', False)

                if services:
                    context = services[0].get('context', '').strip() if services else ''

                    if data.tipo_relatorio == "sp" or is_fachada:
                        # Modo SP ou Fachada: descrição antes das tabelas, 1 tabela por serviço
                        if context:
                            conteudo_final.append({"paragrafo": True})
                            conteudo_final.append({"descricao_texto": context})
                        for s in services:
                            item_obj = s['item']
                            calc = { "largura": s.get('largura', 0.0), "altura": s.get('altura', 0.0), "desconto": s.get('desconto', 0.0), "total": s.get('totalFinal', 0.0) }
                            is_area = item_obj.get('un', '') in ['m²', 'm']
                            conteudo_final.append({"paragrafo": True})
                            conteudo_final.append({"tabela_calculo": { "medicoes": [calc], "is_area": is_area }})
                            conteudo_final.append({"paragrafo": True})
                            conteudo_final.append({"tabela_itens": { "id": item_obj.get('id', ''), "desc": item_obj.get('desc', ''), "total": s.get('totalFinal', 0.0), "un": item_obj.get('un', ''), "totalLabel": s.get("totalLabel", "Total") }})
                    else:
                        # Modo Tradicional: guarda descrição para fim de seção; agrega medidas
                        if context:
                            desc_section = context
                        for s in services:
                            item_obj = s['item']
                            item_id = item_obj.get('id', '')
                            calc = { "largura": s.get('largura', 0.0), "altura": s.get('altura', 0.0), "desconto": s.get('desconto', 0.0), "total": s.get('totalFinal', 0.0) }

                            if item_id not in agregador_tradicional:
                                agregador_tradicional[item_id] = {
                                    "desc": item_obj.get('desc', ''),
                                    "un": item_obj.get('un', ''),
                                    "medicoes": [],
                                    "total_acumulado": 0.0
                                }

                            agregador_tradicional[item_id]["medicoes"].append(calc)
                            agregador_tradicional[item_id]["total_acumulado"] += calc["total"]

        # Flush pra garantir se sobrar algo no fim
        if data.tipo_relatorio == "tradicional":
            flush_tradicional()

        # meta unificado: SP2 usa meta_sp2 (alias de meta_fields vindo do frontend)
        # Outros modos usam meta_fields para substituição de placeholders genéricos
        meta_unified = data.meta_sp2 or data.meta_fields or {}

        if data.tipo_relatorio == "sp":
            from word_utils_sp import inserir_conteudo_sp
            total_images = inserir_conteudo_sp(modelo_path, conteudo_final, output_docx,
                                               selected_description=selected_description,
                                               meta=meta_unified)
            file_logger(f"[OK] Instâncias e imagens inseridas (SP): {total_images}")
        elif data.tipo_relatorio == "sp2":
            nome_pasta_raiz = os.path.basename(pasta_raiz.strip(os.sep))
            output_docx = os.path.join(pasta_saida, f"RELATÓRIO DE VISTORIA - {nome_pasta_raiz} - SP2.docx")
            total_images = inserir_conteudo_sp2(modelo_path, conteudo_final, output_docx,
                                                meta=meta_unified,
                                                selected_description=selected_description)
            file_logger(f"[OK] Instâncias e imagens inseridas (SP2): {total_images}")
        else:
            total_images = generate_report(modelo_path, conteudo_final, output_docx,
                                           logger=file_logger,
                                           selected_description=selected_description,
                                           meta=meta_unified)
        
        return {
            "message": "Relatório gerado com sucesso",
            "total_images": total_images,
            "output_docx": output_docx,
            "log_process": log_process,
            "log_errors": log_errors
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Limpar arquivos temporários de imagens anotadas
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except:
                pass

@app.post("/api/open-folder")
async def open_folder(data: OpenFolderRequest):
    path = data.path.strip()
    if path and os.path.isdir(path):
        try:
            os.startfile(path)
            return {"status": "ok"}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    raise HTTPException(status_code=400, detail="Pasta inválida")

@app.get("/api/download")
async def download_file(path: str = Query(...)):
    """Servindo o arquivo gerado para download direto no navegador."""
    if not path or not os.path.exists(path) or not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
    
    file_name = os.path.basename(path)
    return FileResponse(
        path, 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
        filename=file_name
    )

@app.get("/api/dialog/folder")
async def ask_directory():
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        path = filedialog.askdirectory(parent=root, title="Selecione a pasta")
        root.destroy()
        return {"path": path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=5000)
