import os
import datetime
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image, UnidentifiedImageError

PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

# Constantes de Layout
ALTURA_PADRAO = 10.0  # cm
LARGURA_MAX_3_COL = 5.63  # cm
LARGURA_MAX_2_COL = 7.50  # cm


def cm_to_px(cm):
    return cm * 37.7952755906


def analisar_imagem(caminho_imagem):
    """
    Analisa a imagem e retorna suas características para layout.
    Retorna: (is_vertical, largura_em_10cm)
    """
    try:
        with Image.open(caminho_imagem) as img:
            w, h = img.size
            aspect_ratio = w / h
            largura_em_10cm = ALTURA_PADRAO * aspect_ratio
            
            # Definição de Vertical: Altura > Largura
            # E também não pode ser absurdamente larga (ex: panorâmica vertical)
            is_vertical = h > w
            
            return is_vertical, largura_em_10cm
    except Exception:
        return False, 999.0  # Trata como horizontal/erro em caso de falha


def otimizar_layout(conteudo_original):
    """
    Reorganiza o conteúdo agrupando imagens verticais em tabelas.
    Estratégia Greedy: Tenta agrupar 3, se não der tenta 2, senão 1.
    """
    novo_conteudo = []
    buffer_imagens = []

    def processar_buffer():
        nonlocal buffer_imagens
        while buffer_imagens:
            qtd = len(buffer_imagens)
            grupo_formado = False
            
            # Tenta formar grupo de 3
            if qtd >= 3:
                candidatos = buffer_imagens[:3]
                # Verifica se TODOS cabem na largura de 3 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_3_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 3
                    })
                    buffer_imagens = buffer_imagens[3:]
                    grupo_formado = True
            
            # Se não formou de 3, tenta de 2
            if not grupo_formado and qtd >= 2:
                candidatos = buffer_imagens[:2]
                # Verifica se TODOS cabem na largura de 2 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_2_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 2
                    })
                    buffer_imagens = buffer_imagens[2:]
                    grupo_formado = True
            
            # Se não formou grupo (ou sobrou 1), processa individualmente
            if not grupo_formado:
                img = buffer_imagens.pop(0)
                # Se for "Vertical" mas não coube nos grupos, vira Horizontal (Single)
                # Mantém como imagem normal para ser inserida centralizada
                novo_conteudo.append({"imagem": img['caminho']})

    for item in conteudo_original:
        if isinstance(item, dict) and "imagem" in item:
            caminho = item["imagem"]
            is_vert, larg_10cm = analisar_imagem(caminho)
            
            if is_vert:
                buffer_imagens.append({
                    "caminho": caminho,
                    "largura_10cm": larg_10cm
                })
            else:
                # Se veio uma horizontal, processa o que estava acumulado antes
                processar_buffer()
                novo_conteudo.append(item)
        else:
            # Títulos, quebras, etc. Processa buffer antes.
            processar_buffer()
            novo_conteudo.append(item)

    # Processa qualquer sobra final
    processar_buffer()
    
    return novo_conteudo



def aplicar_estilo(run, size_pt, is_bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = is_bold


def inserir_conteudo(modelo_path, conteudo, output_path):
    doc = Document(modelo_path)
    contador_imagens = 0
    paragrafo_insercao_index = None

    # Localiza a marca de inserção
    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("[AVISO] Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    # Otimiza o layout (agrupamento de imagens)
    conteudo_otimizado = otimizar_layout(conteudo)

    # Processa o conteúdo (reverso para manter ordem de inserção no topo)
    conteudo_invertido = list(reversed(conteudo_otimizado))

    for item in conteudo_invertido:
        # 🟦 Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, True)

        # 🟩 Imagens (Individuais / Horizontais)
        elif isinstance(item, dict) and "imagem" in item:
            imagem_path = item["imagem"]

            if os.path.exists(imagem_path) and os.path.getsize(imagem_path) > 0:
                try:
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size

                        # Altura fixa em 10 cm e largura proporcional
                        altura_desejada_cm = ALTURA_PADRAO
                        proporcao = altura_desejada_cm / (altura_original / 37.7952755906)
                        largura_proporcional_cm = largura_original * proporcao / 37.7952755906

                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        run = p.add_run()
                        run.add_picture(
                            imagem_path,
                            width=Cm(largura_proporcional_cm),
                            height=Cm(altura_desejada_cm)
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1

                except UnidentifiedImageError:
                    msg = f"[ERRO: formato não reconhecido] {imagem_path}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                except Exception as e:
                    msg = f"[ERRO ao inserir imagem: {imagem_path}] {e}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            else:
                msg = f"[ERRO: imagem inválida] {imagem_path}"
                print(msg)
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 🟦 Tabela de Imagens (Agrupadas)
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            
            # Insere parágrafo antes da tabela para separação
            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            # Cria tabela
            tabela = doc.add_table(rows=1, cols=colunas)
            
            # Move a tabela para depois do parágrafo de separação
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            tabela.allow_autofit = False
            
            # Remove bordas (opcional, mas recomendado para layout limpo)
            # Para remover bordas, iteramos sobre as bordas ou usamos estilo padrão sem borda
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)

            # Define largura das colunas (distribuição igual)
            # Largura total página A4 (21cm) - Margens (3cm esq + 2cm dir = 5cm) = 16cm útil
            # Ajuste fino: 16cm / colunas
            largura_coluna = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(largura_coluna)

            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    
                    try:
                        # Insere com altura fixa de 10cm
                        # A largura será proporcional. Como já validamos no otimizar_layout,
                        # sabemos que ela cabe na coluna!
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                    except Exception as e:
                        paragraph.add_run(f"[Erro img: {os.path.basename(img_path)}]")
                        print(f"Erro ao inserir na tabela: {e}")

        # 🟨 Quebra de página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

        # 🟩 Novo tipo: Parágrafo
        elif isinstance(item, dict) and "paragrafo" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')

    # Salva documento final
    doc.save(output_path)
    print(f"[OK] Imagens redimensionadas para 10 cm de altura ({contador_imagens} inseridas).")
    return contador_imagens
