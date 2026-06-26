"""Valida se o DOCX gerado pode ser reaberto pelo python-docx (teste de integridade)."""
import os
import zipfile
from docx import Document

path = r'C:\Users\thiag\TM-MEUS-APPS\0 - NEXT APPS\TM Relatorio teste\output_test\RELATÓRIO FOTOGRÁFICO - ITAMONTE - (Relatório sendo feito) - SP.docx'

print(f"Arquivo existe: {os.path.exists(path)}")
print(f"Tamanho: {os.path.getsize(path) / (1024*1024):.1f} MB")

# Teste 1: Abre como ZIP (DOCX é um ZIP)
try:
    with zipfile.ZipFile(path) as z:
        files = z.namelist()
        print(f"\nArquivos no ZIP: {len(files)}")
        # Verifica se os arquivos essenciais existem
        essenciais = ['word/document.xml', '[Content_Types].xml', 'word/_rels/document.xml.rels']
        for ess in essenciais:
            print(f"  {'OK' if ess in files else 'FALTA'} - {ess}")
except Exception as e:
    print(f"ERRO ao abrir como ZIP: {e}")

# Teste 2: Reabre com python-docx
try:
    doc = Document(path)
    print(f"\npython-docx abriu com sucesso!")
    print(f"  Paragrafos: {len(doc.paragraphs)}")
    print(f"  Tabelas: {len(doc.tables)}")
    
    # Verifica se NÃO há nenhum XML de Shape corrompido
    import re
    with zipfile.ZipFile(path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
        has_shapes = 'wps:wsp' in xml
        has_drawings = 'w:drawing' in xml
        has_shading = 'w:shd' in xml
        print(f"\n  Shapes (wps:wsp): {'SIM ⚠️' if has_shapes else 'NÃO ✅'}")
        print(f"  Drawings (w:drawing): {'SIM ⚠️' if has_drawings else 'NÃO ✅'}")
        print(f"  Shading nativo (w:shd): {'SIM ✅' if has_shading else 'NÃO'}")
        
        # Conta os shading vermelhos
        red_count = xml.count('w:fill="C00000"')
        print(f"  Legendas vermelhas (#C00000): {red_count}")
        
except Exception as e:
    print(f"ERRO ao reabrir com python-docx: {e}")
    import traceback
    traceback.print_exc()
