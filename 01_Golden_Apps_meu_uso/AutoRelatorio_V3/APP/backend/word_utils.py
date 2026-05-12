import os
import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, UnidentifiedImageError

PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

# Constantes de Layout
ALTURA_PADRAO = 10.0  # cm
LARGURA_MAX_3_COL = 5.63  # cm
LARGURA_MAX_2_COL = 7.50  # cm

# Cores das tabelas (MAFFENG V2)
COR_HEADER_CALCULO  = RGBColor(0x00, 0x36, 0x94)   # azul BB
COR_HEADER_ITENS    = RGBColor(0x80, 0x80, 0x80)   # cinza
COR_TOTAL_ROW       = RGBColor(0xAE, 0xAA, 0xAA)   # cinza claro
COR_TEXTO_BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
COR_TEXTO_PRETO     = RGBColor(0x00, 0x00, 0x00)


def cm_to_px(cm):
    return cm * 37.7952755906


def _set_cell_bg(cell, rgb: RGBColor):
    """Define cor de fundo de uma célula via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold=False, font_size=11,
                   color: RGBColor = None, align=WD_PARAGRAPH_ALIGNMENT.CENTER):
    """Escreve texto formatado em uma célula."""
    para = cell.paragraphs[0]
    para.alignment = align
    para.clear()
    run = para.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = color


def inserir_tabela_calculo(doc, paragrafo_ref, medicoes: list, is_area: bool):
    """
    Insere a Tabela de Cálculo MAFFENG antes de `paragrafo_ref`.

    medicoes: lista de dicts com as medidas de cada foto principal.
      Para itens m²:
        { 'largura': 6.5, 'altura': 3.0, 'desconto': 0.0, 'total': 19.5, 'totalLabel': 'Total' }
      Para itens un:
        { 'total': 2.0, 'totalLabel': 'Total' }

    is_area: True → tabela full (larg/alt/desc/total); False → tabela simple (total)
    """
    REFERENCIA = "Foto "   # espaço final intencional para edição manual

    if is_area:
        headers = ["REFERÊNCIA", "LARGURA (m)", "ALTURA (m)", "DESCONTO (m²)", "TOTAL (m²)"]
        col_widths_cm = [3.5, 2.5, 2.5, 2.8, 2.7]
    else:
        headers = ["REFERÊNCIA", "TOTAL (un)"]
        col_widths_cm = [10.0, 4.0]

    n_cols = len(headers)
    n_data = len(medicoes)
    n_rows = 1 + n_data + (1 if n_data > 1 else 0)  # header + dados + total (se múltiplos)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    paragrafo_ref._p.addnext(tbl._tbl)
    tbl.autofit = False
    tbl.style = 'Table Grid'

    # Larguras das colunas
    for col_idx, w in enumerate(col_widths_cm):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    # ── Header (azul) ──
    for col_idx, h in enumerate(headers):
        cell = tbl.cell(0, col_idx)
        _set_cell_bg(cell, COR_HEADER_CALCULO)
        _set_cell_text(cell, h, bold=True, font_size=10, color=COR_TEXTO_BRANCO)

    # ── Linhas de dados ──
    total_acumulado = 0.0
    for row_idx, med in enumerate(medicoes, start=1):
        total = med.get('total', 0.0)
        total_acumulado += total
        cells = tbl.rows[row_idx].cells

        if is_area:
            _set_cell_text(cells[0], REFERENCIA, font_size=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_cell_text(cells[1], f"{med.get('largura', 0.0):.2f}".replace('.', ','), font_size=10)
            _set_cell_text(cells[2], f"{med.get('altura', 0.0):.2f}".replace('.', ','), font_size=10)
            _set_cell_text(cells[3], f"{med.get('desconto', 0.0):.2f}".replace('.', ','), font_size=10)
            _set_cell_text(cells[4], f"{total:.2f}".replace('.', ','), font_size=10, bold=True)
        else:
            _set_cell_text(cells[0], REFERENCIA, font_size=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_cell_text(cells[1], f"{total:.2f}".replace('.', ','), font_size=10, bold=True)

    # ── Linha de Total consolidado (cinza — só se há mais de 1 medida) ──
    if n_data > 1:
        total_row_idx = 1 + n_data
        cells = tbl.rows[total_row_idx].cells
        for cell in cells:
            _set_cell_bg(cell, COR_TOTAL_ROW)

        if is_area:
            _set_cell_text(cells[0], "TOTAL", bold=True, font_size=10, color=COR_TEXTO_BRANCO, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            for i in range(1, 4):
                _set_cell_text(cells[i], "", font_size=10)
            _set_cell_text(cells[4], f"{total_acumulado:.2f}".replace('.', ','), bold=True, font_size=11, color=COR_TEXTO_BRANCO)
        else:
            _set_cell_text(cells[0], "TOTAL", bold=True, font_size=10, color=COR_TEXTO_BRANCO, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_cell_text(cells[1], f"{total_acumulado:.2f}".replace('.', ','), bold=True, font_size=11, color=COR_TEXTO_BRANCO)

    return tbl


def inserir_tabela_itens(doc, paragrafo_ref, item_id: str, item_desc: str,
                         total_final: float, unidade: str, total_label: str):
    """
    Insere a Tabela de Itens MAFFENG (consolidada) antes de `paragrafo_ref`.
    Sempre 1 linha de dados (totais consolidados da seção).

    Colunas: CÓDIGO | DESCRIÇÃO | QUANTIDADE | UNIDADE
    """
    headers = ["CÓDIGO", "DESCRIÇÃO", "QUANTIDADE", "UNIDADE"]
    col_widths_cm = [2.0, 9.0, 2.5, 2.5]

    tbl = doc.add_table(rows=2, cols=4)
    paragrafo_ref._p.addnext(tbl._tbl)
    tbl.autofit = False
    tbl.style = 'Table Grid'

    for col_idx, w in enumerate(col_widths_cm):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    # ── Header "Itens" (cinza) — mesclado ──
    # Insere linha header "Itens" com células mescladas
    header_row = tbl.rows[0]
    for col_idx, h in enumerate(headers):
        cell = header_row.cells[col_idx]
        _set_cell_bg(cell, COR_HEADER_ITENS)
        _set_cell_text(cell, h, bold=True, font_size=10, color=COR_TEXTO_BRANCO)

    # ── Linha de dado ──
    data_cells = tbl.rows[1].cells
    qtd_str = f"{total_final:.2f}".replace('.', ',')
    _set_cell_text(data_cells[0], item_id, font_size=10)
    _set_cell_text(data_cells[1], item_desc, font_size=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    _set_cell_text(data_cells[2], qtd_str, font_size=10, bold=True)
    _set_cell_text(data_cells[3], unidade, font_size=10)

    return tbl


def analisar_imagem(caminho_imagem):
    """
    Analisa a imagem e retorna suas características para layout.
    Retorna: (is_vertical, largura_em_10cm)
    """
    try:
        with Image.open(caminho_imagem) as img:
            w, h = img.size
            aspect_ratio = w / h
            largura_em_10cm = ALTURA_PADRAO * aspect_ratio
            
            # Definição de Vertical: Altura > Largura
            # E também não pode ser absurdamente larga (ex: panorâmica vertical)
            is_vertical = h > w
            
            return is_vertical, largura_em_10cm
    except Exception:
        return False, 999.0  # Trata como horizontal/erro em caso de falha


def otimizar_layout(conteudo_original):
    """
    Reorganiza o conteúdo agrupando imagens verticais em tabelas.
    Estratégia Greedy: Tenta agrupar 3, se não der tenta 2, senão 1.
    """
    novo_conteudo = []
    buffer_imagens = []

    def processar_buffer():
        nonlocal buffer_imagens
        while buffer_imagens:
            qtd = len(buffer_imagens)
            grupo_formado = False
            
            # Tenta formar grupo de 3
            if qtd >= 3:
                candidatos = buffer_imagens[:3]
                # Verifica se TODOS cabem na largura de 3 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_3_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 3
                    })
                    buffer_imagens = buffer_imagens[3:]
                    grupo_formado = True
            
            # Se não formou de 3, tenta de 2
            if not grupo_formado and qtd >= 2:
                candidatos = buffer_imagens[:2]
                # Verifica se TODOS cabem na largura de 2 colunas
                if all(img['largura_10cm'] <= LARGURA_MAX_2_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 2
                    })
                    buffer_imagens = buffer_imagens[2:]
                    grupo_formado = True
            
            # Se não formou grupo (ou sobrou 1), processa individualmente
            if not grupo_formado:
                img = buffer_imagens.pop(0)
                # Se for "Vertical" mas não coube nos grupos, vira Horizontal (Single)
                # Mantém como imagem normal para ser inserida centralizada
                novo_conteudo.append({"imagem": img['caminho']})

    for item in conteudo_original:
        if isinstance(item, dict) and "imagem" in item:
            caminho = item["imagem"]
            is_vert, larg_10cm = analisar_imagem(caminho)
            
            if is_vert:
                buffer_imagens.append({
                    "caminho": caminho,
                    "largura_10cm": larg_10cm
                })
            else:
                # Se veio uma horizontal, processa o que estava acumulado antes
                processar_buffer()
                novo_conteudo.append(item)
        else:
            # Títulos, quebras, etc. Processa buffer antes.
            processar_buffer()
            novo_conteudo.append(item)

    # Processa qualquer sobra final
    processar_buffer()
    
    return novo_conteudo



def aplicar_estilo(run, size_pt, is_bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = is_bold


def inserir_conteudo(modelo_path, conteudo, output_path, selected_description=None):
    doc = Document(modelo_path)

    # Substituição do placeholder {Desc_here} no documento inteiro (parágrafos e tabelas)
    if selected_description:
        # Substituir nos parágrafos
        for p in doc.paragraphs:
            if "{Desc_here}" in p.text:
                p.text = p.text.replace("{Desc_here}", selected_description)
        
        # Substituir nas tabelas (caso o placeholder esteja dentro de uma célula)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{Desc_here}" in p.text:
                            p.text = p.text.replace("{Desc_here}", selected_description)

    contador_imagens = 0
    paragrafo_insercao_index = None

    # Localiza a marca de inserção
    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("[AVISO] Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    # Otimiza o layout (agrupamento de imagens)
    conteudo_otimizado = otimizar_layout(conteudo)

    # Processa o conteúdo (reverso para manter ordem de inserção no topo)
    conteudo_invertido = list(reversed(conteudo_otimizado))

    for item in conteudo_invertido:
        # 🟦 Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, True)

        # 🟩 Imagens (Individuais / Horizontais)
        elif isinstance(item, dict) and "imagem" in item:
            imagem_path = item["imagem"]

            if os.path.exists(imagem_path) and os.path.getsize(imagem_path) > 0:
                try:
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size

                        # Altura fixa em 10 cm e largura proporcional
                        altura_desejada_cm = ALTURA_PADRAO
                        proporcao = altura_desejada_cm / (altura_original / 37.7952755906)
                        largura_proporcional_cm = largura_original * proporcao / 37.7952755906

                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        run = p.add_run()
                        run.add_picture(
                            imagem_path,
                            width=Cm(largura_proporcional_cm),
                            height=Cm(altura_desejada_cm)
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1

                except UnidentifiedImageError:
                    msg = f"[ERRO: formato não reconhecido] {imagem_path}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                except Exception as e:
                    msg = f"[ERRO ao inserir imagem: {imagem_path}] {e}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            else:
                msg = f"[ERRO: imagem inválida] {imagem_path}"
                print(msg)
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 🟦 Tabela de Imagens (Agrupadas)
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            
            # Insere parágrafo antes da tabela para separação
            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            # Cria tabela
            tabela = doc.add_table(rows=1, cols=colunas)
            
            # Move a tabela para depois do parágrafo de separação
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            tabela.allow_autofit = False
            
            # Remove bordas (opcional, mas recomendado para layout limpo)
            # Para remover bordas, iteramos sobre as bordas ou usamos estilo padrão sem borda
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)

            # Define largura das colunas (distribuição igual)
            # Largura total página A4 (21cm) - Margens (3cm esq + 2cm dir = 5cm) = 16cm útil
            # Ajuste fino: 16cm / colunas
            largura_coluna = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(largura_coluna)

            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    
                    try:
                        # Insere com altura fixa de 10cm
                        # A largura será proporcional. Como já validamos no otimizar_layout,
                        # sabemos que ela cabe na coluna!
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                    except Exception as e:
                        paragraph.add_run(f"[Erro img: {os.path.basename(img_path)}]")
                        print(f"Erro ao inserir na tabela: {e}")

        # 🟨 Quebra de página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].add_run().add_break(WD_BREAK.PAGE)

        # 🟩 Novo tipo: Parágrafo em branco
        elif isinstance(item, dict) and "paragrafo" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')

        # 📊 Tabela de Cálculo MAFFENG (nova — Fase 3)
        elif isinstance(item, dict) and "tabela_calculo" in item:
            dados = item["tabela_calculo"]   # { medicoes: [...], is_area: bool }
            p_ref = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            inserir_tabela_calculo(
                doc, p_ref,
                medicoes=dados.get("medicoes", []),
                is_area=dados.get("is_area", True)
            )

        # 📋 Tabela de Itens MAFFENG (nova — Fase 3)
        elif isinstance(item, dict) and "tabela_itens" in item:
            dados = item["tabela_itens"]
            p_ref = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            inserir_tabela_itens(
                doc, p_ref,
                item_id=dados.get("id", ""),
                item_desc=dados.get("desc", ""),
                total_final=dados.get("total", 0.0),
                unidade=dados.get("un", ""),
                total_label=dados.get("totalLabel", "Total")
            )

        # 📝 Descrição Técnica (texto livre da Observação do modal)
        elif isinstance(item, dict) and "descricao_texto" in item:
            texto = item.get("descricao_texto", "").strip()
            if texto:
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                run = p.add_run(texto)
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_before = Pt(2)  # 40 twips
                p.paragraph_format.space_after = Pt(2)

    # Salva documento final
    doc.save(output_path)
    print(f"[OK] Imagens redimensionadas para 10 cm de altura ({contador_imagens} inseridas).")
    return contador_imagens

