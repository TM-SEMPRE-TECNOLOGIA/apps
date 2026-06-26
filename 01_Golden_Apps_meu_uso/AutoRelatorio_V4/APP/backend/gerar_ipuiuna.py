"""
Gerador direto para IPUIUNA — contrato 2057 modo Tradicional
Motor extraido do Arquivo Morto (NX Relatorios SP - Mais atualizado)
- Imagens a 10cm de altura, proporcional
- Verticais agrupadas em tabelas 2 ou 3 colunas (greedy)
- Vista ampla e Detalhes: sem heading, texto normal em negrito
- Ordem correta: insercao invertida no insert_paragraph_before
- Quebra de pagina: WD_BREAK.PAGE (sem bug)
"""

import os
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image, UnidentifiedImageError

# ── Paths ─────────────────────────────────────────────────────────────────────

PASTA_RAIZ = r'C:\Users\thiag\Desktop\000 - Minha Demanda\2 - Em andamento\IPUIUNA -'
MODELO     = r'C:\Users\thiag\Desktop\TM-MEUS-APPS\01_Golden_Apps_meu_uso\AutoRelatorio_V4\APP\backend\templates\MODELO - 2057 - VARGINHA - ATUALIZADO.docx'
OUTPUT     = r'C:\Users\thiag\Desktop\000 - Minha Demanda\2 - Em andamento\IPUIUNA -\- Envio portal\RELATÓRIO FOTOGRÁFICO - IPUIUNA - LEVANTAMENTO PREVENTIVO.docx'

# ── Constantes ────────────────────────────────────────────────────────────────

PASTAS_IGNORAR     = {'- envio portal', 'envio portal', 'output', 'outputs'}
ORDEM_PASTAS       = ['- Área externa', '- Área interna', '- Segundo piso']
PASTAS_TEXTO_NORMAL = ['- Detalhes', '- Vista ampla']   # nao viram heading

ALTURA_PADRAO    = 10.0   # cm
LARGURA_MAX_2COL =  7.50  # cm — maximo para 2 verticais lado a lado
LARGURA_MAX_3COL =  5.63  # cm — maximo para 3 verticais lado a lado

PX_PER_CM = 37.7952755906

# ── Scanner ───────────────────────────────────────────────────────────────────

def folder_sort_key(name):
    nl = name.lower()
    if 'vista ampla' in nl:
        return (0, 0, nl)
    m = re.match(r'^(\d+)(.*)', name)
    if m:
        return (1, int(m.group(1)), m.group(2))
    if 'detalhes' in nl:
        return (3, 0, nl)
    return (2, 0, nl)


def scan(pasta_raiz):
    conteudo = []

    for root_dir, dirs, files in os.walk(pasta_raiz, topdown=True):
        root_dir = os.fsdecode(root_dir)
        files    = [os.fsdecode(f) for f in files]

        # Filtra pastas ignoradas
        dirs[:] = [os.fsdecode(d) for d in dirs
                   if d.lower() not in PASTAS_IGNORAR]

        # Ordena subpastas
        if root_dir == pasta_raiz:
            dirs.sort(key=lambda x: (
                0 if x == '- Vista ampla' else 1,
                ORDEM_PASTAS.index(x) if x in ORDEM_PASTAS else len(ORDEM_PASTAS),
                folder_sort_key(x),
            ))
        else:
            dirs.sort(key=folder_sort_key)

        path_parts = os.path.relpath(root_dir, pasta_raiz).split(os.sep)
        nome  = path_parts[-1]
        nivel = len(path_parts)

        imgs = sorted(
            [os.path.join(root_dir, f) for f in files
             if f.lower().endswith(('.jpg', '.jpeg', '.png'))],
            key=lambda x: os.path.getctime(x) if os.path.exists(x) else x
        )

        # Raiz: apenas fachada no topo, sem titulo de secao
        if nome == '.':
            fachada = [i for i in imgs
                       if 'fachada' in os.path.basename(i).lower()]
            if fachada:
                conteudo.append('- Fachada')
                for f in fachada:
                    conteudo.append({'imagem': f})
                conteudo.append({'quebra_pagina': True})
            continue

        # Titulo da pasta (prefixo » define nivel)
        prefixos = {1: '', 2: '»', 3: '»»'}
        conteudo.append(f"{prefixos.get(nivel, '»»»')}{nome}")

        for img in imgs:
            conteudo.append({'imagem': img})

        conteudo.append({'quebra_pagina': True})

    return conteudo

# ── Layout optimizer (motor arquivo morto — greedy 3/2/1) ────────────────────

def analisar_imagem(caminho):
    try:
        with Image.open(caminho) as img:
            w, h = img.size
            largura_em_10cm = ALTURA_PADRAO * (w / h)
            is_vertical = h > w
            return is_vertical, largura_em_10cm
    except Exception:
        return False, 999.0


def otimizar_layout(conteudo_original):
    novo_conteudo = []
    buffer_imgs   = []

    def processar_buffer():
        while buffer_imgs:
            qtd = len(buffer_imgs)
            grupo = False

            if qtd >= 3:
                cands = buffer_imgs[:3]
                if all(img['larg'] <= LARGURA_MAX_3COL for img in cands):
                    novo_conteudo.append({
                        'tabela_imagens': [i['path'] for i in cands],
                        'colunas': 3
                    })
                    del buffer_imgs[:3]
                    grupo = True

            if not grupo and qtd >= 2:
                cands = buffer_imgs[:2]
                if all(img['larg'] <= LARGURA_MAX_2COL for img in cands):
                    novo_conteudo.append({
                        'tabela_imagens': [i['path'] for i in cands],
                        'colunas': 2
                    })
                    del buffer_imgs[:2]
                    grupo = True

            if not grupo:
                img = buffer_imgs.pop(0)
                novo_conteudo.append({'imagem': img['path']})

    for item in conteudo_original:
        if isinstance(item, dict) and 'imagem' in item:
            caminho = item['imagem']
            is_vert, larg = analisar_imagem(caminho)
            if is_vert:
                buffer_imgs.append({'path': caminho, 'larg': larg})
            else:
                processar_buffer()
                novo_conteudo.append(item)
        else:
            processar_buffer()
            novo_conteudo.append(item)

    processar_buffer()
    return novo_conteudo

# ── Gerador ───────────────────────────────────────────────────────────────────

def aplicar_estilo(run, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = 'Arial'


def gerar(modelo, conteudo, output):
    doc = Document(modelo)

    # Substitui placeholders (em branco para preenchimento manual)
    meta = {
        'nr_os': '',
        'ag_cod': '',
        'ag_nome': 'IPUIUNA',
        'endereco': '',
        'responsavel_dependencia': '',
        'dt_atend': '',
        'dt_elab': '',
    }
    for para in doc.paragraphs:
        for k, v in meta.items():
            if f'{{{{{k}}}}}' in para.text:
                para.text = para.text.replace(f'{{{{{k}}}}}', v)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for k, v in meta.items():
                        if f'{{{{{k}}}}}' in para.text:
                            para.text = para.text.replace(f'{{{{{k}}}}}', v)

    # Localiza ancora {{start_here}}
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if '{{start_here}}' in para.text:
            para.text = ''
            anchor_idx = i
            break

    if anchor_idx is None:
        print('ERRO: {{start_here}} nao encontrado!')
        return 0

    # Otimiza layout (agrupamento de verticais)
    conteudo_otimizado = otimizar_layout(conteudo)

    # CHAVE: inverte para que insert_paragraph_before mantenha ordem correta
    conteudo_invertido = list(reversed(conteudo_otimizado))

    total_imgs = 0

    for item in conteudo_invertido:
        anchor = doc.paragraphs[anchor_idx]

        try:
            # ── Titulo ──────────────────────────────────────────────────────
            if isinstance(item, str):
                nivel  = item.count('»')
                titulo = item.replace('»', '').strip() + ':'
                p = anchor.insert_paragraph_before('')
                run = p.add_run(titulo)

                if any(t in titulo for t in PASTAS_TEXTO_NORMAL):
                    # Vista ampla / Detalhes: texto normal negrito, sem heading
                    aplicar_estilo(run, 11, True)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                elif nivel == 0:
                    p.style = 'Heading 1'
                elif nivel == 1:
                    p.style = 'Heading 2'
                elif nivel == 2:
                    p.style = 'Heading 3'
                else:
                    aplicar_estilo(run, 11, True)

            # ── Imagem individual ────────────────────────────────────────────
            elif isinstance(item, dict) and 'imagem' in item:
                caminho = item['imagem']
                if not os.path.exists(caminho) or os.path.getsize(caminho) == 0:
                    print(f'  [AVISO] Imagem invalida: {caminho}')
                    continue
                try:
                    with Image.open(caminho) as img:
                        w_px, h_px = img.size
                        prop    = ALTURA_PADRAO / (h_px / PX_PER_CM)
                        w_cm    = (w_px * prop) / PX_PER_CM
                    p = anchor.insert_paragraph_before('')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.add_run().add_picture(caminho, width=Cm(w_cm), height=Cm(ALTURA_PADRAO))
                    total_imgs += 1
                except UnidentifiedImageError:
                    print(f'  [ERRO formato] {os.path.basename(caminho)}')
                except Exception as e:
                    print(f'  [ERRO img] {os.path.basename(caminho)}: {e}')

            # ── Tabela de imagens (2 ou 3 colunas) ──────────────────────────
            elif isinstance(item, dict) and 'tabela_imagens' in item:
                imagens  = item['tabela_imagens']
                colunas  = item['colunas']

                p_sep = anchor.insert_paragraph_before('')
                tabela = doc.add_table(rows=1, cols=colunas)
                p_sep._p.addnext(tabela._tbl)
                tabela.autofit = False

                # Remove bordas
                for border in tabela._tbl.tblPr.xpath('./w:tblBorders'):
                    border.getparent().remove(border)

                largura_col = 16.0 / colunas
                for col in tabela.columns:
                    col.width = Cm(largura_col)

                for idx, img_path in enumerate(imagens):
                    if idx >= colunas:
                        break
                    cell = tabela.cell(0, idx)
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    try:
                        para.add_run().add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        total_imgs += 1
                    except Exception as e:
                        para.add_run(f'[Erro: {os.path.basename(img_path)}]')
                        print(f'  [ERRO tabela] {os.path.basename(img_path)}: {e}')

            # ── Quebra de pagina ─────────────────────────────────────────────
            elif isinstance(item, dict) and 'quebra_pagina' in item:
                anchor.insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

        except Exception as e:
            print(f'  [ERRO item] {e}')
            continue

    os.makedirs(os.path.dirname(output), exist_ok=True)
    doc.save(output)
    return total_imgs

# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Escaneando pastas...')
    conteudo = scan(PASTA_RAIZ)
    n_imgs   = sum(1 for x in conteudo if isinstance(x, dict) and 'imagem' in x)
    n_tits   = sum(1 for x in conteudo if isinstance(x, str))
    print(f'  {n_tits} titulos, {n_imgs} imagens esperadas')

    print('Gerando DOCX...')
    total = gerar(MODELO, conteudo, OUTPUT)
    print(f'  {total}/{n_imgs} imagens inseridas')
    print(f'Salvo em: {OUTPUT}')
