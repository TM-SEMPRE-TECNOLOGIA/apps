#!/usr/bin/env python
"""Insere os 7 placeholders dinâmicos em todos os templates"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

TEMPLATES_DIR = Path(__file__).parent / "templates"

PLACEHOLDERS = [
    'nr_os',
    'data_elaboracao',
    'data_atendimento',
    'agencia_codigo',
    'agencia_nome',
    'endereco',
    'responsavel_dependencia'
]

PLACEHOLDER_TEXT = """
[PLACEHOLDERS DINÂMICOS v3.2]
Este documento contém os seguintes placeholders que serão substituídos:
- {{nr_os}} - Número da Ordem de Serviço
- {{data_elaboracao}} - Data de Elaboração
- {{data_atendimento}} - Data do Atendimento
- {{agencia_codigo}} - Código da Agência
- {{agencia_nome}} - Nome da Agência
- {{endereco}} - Endereço da Dependência
- {{responsavel_dependencia}} - Responsável da Dependência
"""

def insert_placeholders_in_template(template_path: Path):
    """Insere os placeholders em um template"""
    try:
        print(f"[PROCESSAR] {template_path.name}...")
        doc = Document(template_path)

        # Procurar por {{start_here}}
        placeholder_para = None
        for i, para in enumerate(doc.paragraphs):
            if "{{start_here}}" in para.text:
                placeholder_para = para
                break

        if placeholder_para is None:
            print(f"  [ERRO] Marcador {{{{start_here}}}} não encontrado")
            return False

        # Limpar parágrafo
        placeholder_para.text = ""

        # Adicionar placeholder section com informações
        p = placeholder_para.insert_paragraph_before()
        run = p.add_run(PLACEHOLDER_TEXT)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Adicionar quebra de página antes do conteúdo antigo
        p_break = placeholder_para.insert_paragraph_before()
        pPr = p_break._element.get_or_add_pPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr.append(OxmlElement('w:pageBreakBefore'))

        # Manter {{start_here}} para compatibilidade com sistema antigo
        placeholder_para.text = "{{start_here}}"

        # Salvar
        doc.save(template_path)
        print(f"  [OK] Placeholders inseridos")
        return True

    except Exception as e:
        print(f"  [ERRO] {e}")
        return False

print("[SCRIPT] Inserindo placeholders dinâmicos em templates...\n")

if not TEMPLATES_DIR.exists():
    print(f"[ERRO] Pasta de templates não encontrada: {TEMPLATES_DIR}")
    exit(1)

templates = list(TEMPLATES_DIR.glob("*.docx"))
print(f"[INFO] {len(templates)} templates encontrados\n")

success_count = 0
for template_path in sorted(templates):
    if insert_placeholders_in_template(template_path):
        success_count += 1

print(f"\n[RESULTADO] {success_count}/{len(templates)} templates processados com sucesso")

if success_count == len(templates):
    print("[SUCCESS] Todos os templates atualizados!")
else:
    print(f"[AVISO] {len(templates) - success_count} template(s) falharam")
