import sys
import os
from docx import Document

# Adiciona o diretório de referências da skill v2 ao path
sys.path.insert(0, r"C:\Users\thiag\Desktop\TM-MEUS-APPS\Meus Plugins e Skills\relatorio-preventivo-v2\skills\relatorio-preventivo\references")

docx_path = r"C:\Users\thiag\Downloads\RELATÓRIO FOTOGRÁFICO - VILA TIBERIO - LEVANTAMENTO PREVENTIVO(1).docx"

print(f"Lendo arquivo: {docx_path}")

try:
    with open(docx_path, 'rb') as docx_file:
        doc = Document(docx_file)
    print("Sucesso ao carregar o Documento com stream binário!")
    print(f"Número de tabelas encontradas no documento: {len(doc.tables)}")
    
    from parse_report import iter_block_items, raw_row, detect_htype, parse_float
    
    def parse_docx_stream(stream):
        doc = Document(stream)
        occurrences = []
        
        for block in iter_block_items(doc):
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            
            if isinstance(block, Paragraph):
                text = block.text.strip().lower()
                if "memorial de cálculo e itens" in text or "resumo geral" in text or "memorial final" in text:
                    break
                continue

            if isinstance(block, Table):
                table = block
                rows = table.rows
                if len(rows) < 4:
                    continue

                r0 = raw_row(rows[0])
                r1 = raw_row(rows[1])
                r2 = raw_row(rows[2])

                if not r0 or 'Itens' not in r0[0]:
                    continue

                code = r1[0].strip()
                desc = r1[1].strip() if len(r1) > 1 else ''
                
                headers = r2
                htype = detect_htype(headers)
                last_header = headers[-1] if headers else 'Total (m²)'

                data_rows = []
                unit = ''
                for ri in range(3, len(rows) - 1):
                    r = raw_row(rows[ri])
                    if not r or not r[0].strip():
                        continue
                    if r[0].strip().lower().startswith('total'):
                        continue

                    if htype == 'full':
                        data_rows.append({
                            'foto': r[0] if len(r) > 0 else '',
                            'comp': r[1] if len(r) > 1 else '',
                            'alt':  r[2] if len(r) > 2 else '',
                            'desc': r[3] if len(r) > 3 else '',
                            'sub':  r[4] if len(r) > 4 else '',
                            'tot':  r[5] if len(r) > 5 else '',
                        })
                    elif htype == 'nodiscount':
                        data_rows.append({
                            'foto': r[0] if len(r) > 0 else '',
                            'comp': r[1] if len(r) > 1 else '',
                            'alt':  r[2] if len(r) > 2 else '',
                            'tot':  r[3] if len(r) > 3 else '',
                        })
                    else:  # unit
                        un = r[2] if len(r) > 2 else ''
                        if un:
                            unit = un
                        data_rows.append({
                            'foto': r[0] if len(r) > 0 else '',
                            'qty':  r[1] if len(r) > 1 else '',
                            'unit': un,
                        })

                total_row = raw_row(rows[-1])
                total_label = total_row[0].strip() if total_row else 'Total'
                total_value = 0.0

                if htype == 'unit':
                    if len(total_row) >= 3:
                        total_value = parse_float(total_row[-2])
                        unit = total_row[-1] or unit
                    elif len(total_row) == 2:
                        total_value = parse_float(total_row[-1])
                    if not unit:
                        for r in data_rows:
                            if r.get('unit'):
                                unit = r['unit']
                                break
                else:
                    total_value = parse_float(total_row[-1]) if total_row else 0.0

                occurrences.append({
                    'code':        code,
                    'desc':        desc,
                    'htype':       htype,
                    'headers':     headers,
                    'last_header': last_header,
                    'rows':        data_rows,
                    'total_value': total_value,
                    'total_label': total_label,
                    'unit':        unit,
                })
        return occurrences

    with open(docx_path, 'rb') as docx_file:
        occurrences = parse_docx_stream(docx_file)
        
    print(f"\nSucesso! Total de tabelas filtradas: {len(occurrences)}")
    
    types = {}
    for occ in occurrences:
        ht = occ['htype']
        types[ht] = types.get(ht, 0) + 1
    
    print("\nTipos de tabelas encontrados:")
    for t, count in types.items():
        print(f"  - {t}: {count} ocorrência(s)")
        
    print("\nExemplo das primeiras 2 ocorrências:")
    for idx, occ in enumerate(occurrences[:2]):
        print("-" * 50)
        print(f"Ocorrência {idx+1}:")
        print(f"  Código: {occ['code']}")
        print(f"  Descrição: {occ['desc']}")
        print(f"  Tipo: {occ['htype']}")
        print(f"  Headers: {occ['headers']}")
        print(f"  Total Declarado: {occ['total_value']} {occ['unit']}")
        
except Exception as e:
    import traceback
    print("Erro ao ler o documento:")
    traceback.print_exc()
