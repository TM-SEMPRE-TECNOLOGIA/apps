import sys
import os
import io
import re
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

sys.path.insert(0, r"C:\Users\thiag\Desktop\TM-MEUS-APPS\Meus Plugins e Skills\relatorio-preventivo-v2\skills\relatorio-preventivo\references")
from parse_report import iter_block_items, raw_row, parse_float

# ── helpers ─────────────────────────────────────────────────────
def fmt(v):
    return f"{v:.2f}".replace('.', ',')

# ── Classificação de tabela ─────────────────────────────────────
# Existem 3 tipos REAIS no documento:
#   full        → 6 colunas: Foto | Comp | Altura | Desconto | Subtotal | Total
#   nodiscount  → 4 colunas: Foto | Comp | Altura | Total  (sem desconto)
#   unit        → 3 colunas: Foto | Quantitativo | Unidade
#                 2 colunas: Foto | Total  (variante simplificada de unit)
#
# Para o MEMORIAL FINAL, o output sempre será:
#   full/nodiscount → renderizados com 6 colunas (nodiscount recebe Desconto=0,00 e Subtotal=Total)
#   unit            → renderizados com 3 colunas (Foto | Quantitativo | Unidade)

def classify_table(headers):
    """Classifica a tabela pelo conteúdo dos cabeçalhos."""
    h_lower = [str(h).lower().strip() for h in headers]
    ncols = len(headers)
    
    # 2 colunas: Foto | Total → unit
    if ncols == 2:
        return 'unit'
    
    # 3 colunas com "quantitativo" → unit
    if ncols == 3 and 'quantitativo' in h_lower[1]:
        return 'unit'
    
    # 6 colunas com "desconto" → full
    if ncols >= 6:
        return 'full'
    
    # 4 colunas sem desconto → nodiscount
    if ncols == 4:
        return 'nodiscount'
    
    # fallback
    return 'full'

def parse_data_row_full(r):
    """Extrai dados de uma linha de tabela com 6 colunas."""
    return {
        'foto': r[0] if len(r) > 0 else '',
        'comp': r[1] if len(r) > 1 else '',
        'alt':  r[2] if len(r) > 2 else '',
        'desc': r[3] if len(r) > 3 else '0,00',
        'sub':  r[4] if len(r) > 4 else '',
        'tot':  r[5] if len(r) > 5 else '',
    }

def parse_data_row_nodiscount(r):
    """Extrai dados de uma linha de tabela com 4 colunas (sem desconto).
    Converte para o formato full com Desconto=0,00 e Subtotal=Total."""
    tot = r[3] if len(r) > 3 else ''
    return {
        'foto': r[0] if len(r) > 0 else '',
        'comp': r[1] if len(r) > 1 else '',
        'alt':  r[2] if len(r) > 2 else '',
        'desc': '0,00',
        'sub':  tot,
        'tot':  tot,
    }

def parse_data_row_unit(r):
    """Extrai dados de uma linha de tabela unitária (2 ou 3 colunas)."""
    return {
        'foto': r[0] if len(r) > 0 else '',
        'qty':  r[1] if len(r) > 1 else '',
        'unit': r[2] if len(r) > 2 else '',
    }

# ── Parser principal ────────────────────────────────────────────
def parse_docx_stream(stream):
    doc = Document(stream)
    occurrences = []
    
    for block in iter_block_items(doc):
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        
        if isinstance(block, Paragraph):
            text = block.text.strip().lower()
            if any(k in text for k in ["memorial de cálculo e itens", "resumo geral", "memorial final"]):
                break
            continue

        if isinstance(block, Table):
            table = block
            rows = table.rows
            if len(rows) < 4:
                continue

            # Encontra o índice da linha de cabeçalhos (que começa com "Foto")
            header_idx = -1
            for i in range(min(4, len(rows))):
                r = raw_row(rows[i])
                if r and len(r) > 0 and 'foto' in str(r[0]).lower().strip():
                    header_idx = i
                    break
            
            if header_idx < 1 or len(rows) <= header_idx + 1:
                continue

            # Linha de Código + Descrição é a linha imediatamente anterior aos cabeçalhos
            r1 = raw_row(rows[header_idx - 1])
            code = r1[0].strip()
            desc = r1[1].strip() if len(r1) > 1 else ''
            
            # Remove parênteses do código se existirem (ex: "(17.7)")
            code = re.sub(r'^\((.*?)\)', r'\1', code).strip()

            # Células mescladas: código e descrição vêm juntos
            if desc == code or not desc:
                parts = code.split(' - ', 1)
                if len(parts) == 2:
                    code = parts[0].strip()
                    desc = parts[1].strip()

            # Cabeçalhos
            headers = raw_row(rows[header_idx])
            htype = classify_table(headers)
            last_header = headers[-1] if headers else 'Total (m²)'

            # Linhas de dados
            data_rows = []
            unit = ''
            for ri in range(header_idx + 1, len(rows) - 1):
                r = raw_row(rows[ri])
                if not r or not r[0].strip():
                    continue
                if r[0].strip().lower().startswith('total'):
                    continue

                if htype == 'full':
                    data_rows.append(parse_data_row_full(r))
                elif htype == 'nodiscount':
                    data_rows.append(parse_data_row_nodiscount(r))
                else:  # unit
                    row_data = parse_data_row_unit(r)
                    if row_data['unit']:
                        unit = row_data['unit']
                    data_rows.append(row_data)

            # Linha Total
            total_row = raw_row(rows[-1])
            total_label = total_row[0].strip() if total_row else 'Total'
            total_value = 0.0

            if htype == 'unit':
                if len(total_row) >= 3:
                    total_value = parse_float(total_row[-2])
                    u = total_row[-1]
                    if u:
                        unit = u
                elif len(total_row) == 2:
                    val_str = str(total_row[-1])
                    total_value = parse_float(val_str)
                    letters = ''.join([c for c in val_str if c.isalpha()]).strip().upper()
                    if letters:
                        unit = letters
                if not unit:
                    for r in data_rows:
                        if r.get('unit'):
                            unit = r['unit']
                            break
            else:
                total_value = parse_float(total_row[-1]) if total_row else 0.0

            # Para o memorial, nodiscount vira full (já convertemos as linhas acima)
            output_htype = 'unit' if htype == 'unit' else 'full'

            occurrences.append({
                'code':        code,
                'desc':        desc,
                'htype':       output_htype,
                'headers':     headers,
                'last_header': last_header,
                'rows':        data_rows,
                'total_value': total_value,
                'total_label': total_label,
                'unit':        unit,
            })
    return occurrences

# ── Consolidação ────────────────────────────────────────────────
def consolidate(occurrences):
    consolidated = OrderedDict()
    for occ in occurrences:
        c = occ['code']
        if c not in consolidated:
            consolidated[c] = {
                'code':        c,
                'desc':        occ['desc'],
                'htype':       occ['htype'], 
                'last_header': occ['last_header'],
                'rows':        [],
                'total':       0.0,
                'unit':        occ['unit'],
                'total_label': 'Total',
                'labels':      set(),
            }
        
        # Se misturar full e unit do mesmo código, prevalece full
        if occ['htype'] == 'full' and consolidated[c]['htype'] == 'unit':
            consolidated[c]['htype'] = 'full'
            
        consolidated[c]['rows'].extend(occ['rows'])
        consolidated[c]['total'] += occ['total_value']
        consolidated[c]['labels'].add(occ['total_label'])
        if occ['unit'] and not consolidated[c]['unit']:
            consolidated[c]['unit'] = occ['unit']

    for c, d in consolidated.items():
        # No consolidado, o label é sempre "Total" porque os valores
        # individuais já incluem os multiplicadores (X2, X3, etc.)
        d['total_label'] = 'Total'
    return consolidated

# ── Export Excel ────────────────────────────────────────────────
def thin_border():
    s = Side(style='thin', color='CCCCCC')
    return Border(left=s, right=s, top=s, bottom=s)

def header_cell(ws, row, col, text):
    c = ws.cell(row=row, column=col, value=text)
    c.font = Font(bold=True, name='Calibri', size=10, color='FFFFFF')
    c.fill = PatternFill('solid', start_color='1F4E79')
    c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    c.border = thin_border()

def data_cell(ws, row, col, value, align='left', bold=False, bg='FFFFFF'):
    c = ws.cell(row=row, column=col, value=value)
    c.font = Font(name='Calibri', size=9, bold=bold)
    c.fill = PatternFill('solid', start_color=bg)
    c.alignment = Alignment(horizontal=align, vertical='center', wrap_text=True)
    c.border = thin_border()

def build_sheet(ws, title, items_list, qty_col_header='Quantidade'):
    ws.row_dimensions[1].height = 28
    ws.merge_cells('A1:D1')
    t = ws.cell(row=1, column=1, value=title)
    t.font = Font(bold=True, name='Calibri', size=12, color='FFFFFF')
    t.fill = PatternFill('solid', start_color='1F4E79')
    t.alignment = Alignment(horizontal='center', vertical='center')

    ws.row_dimensions[2].height = 20
    for ci, h in enumerate(['Codigo', 'Descricao', qty_col_header, 'Unidade'], 1):
        header_cell(ws, 2, ci, h)

    for ri, item in enumerate(items_list, 3):
        bg = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        data_cell(ws, ri, 1, item['code'], align='center', bg=bg)
        data_cell(ws, ri, 2, item['desc'], bg=bg)
        data_cell(ws, ri, 3, item['qty_str'], align='center', bg=bg)
        data_cell(ws, ri, 4, item['unit'], align='center', bg=bg)
        ws.row_dimensions[ri].height = 28

    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 70
    ws.column_dimensions['C'].width = 16
    ws.column_dimensions['D'].width = 12

def export_excel(occurrences, consolidated, out_path):
    all_items = []
    for occ in occurrences:
        all_items.append({
            'code': occ['code'],
            'desc': occ['desc'],
            'qty_str': fmt(occ['total_value']),
            'unit': occ['unit']
        })

    consol_list = []
    for code, d in consolidated.items():
        consol_list.append({
            'code': code,
            'desc': d['desc'],
            'qty_str': fmt(d['total']),
            'unit': d['unit']
        })

    wb = Workbook()
    ws1 = wb.active
    ws1.title = 'Lista Completa'
    build_sheet(ws1, 'LISTA DE ITENS - ORDEM DO RELATORIO', all_items)

    ws2 = wb.create_sheet('Consolidado')
    build_sheet(ws2, 'LISTA CONSOLIDADA - QUANTIDADES SOMADAS', consol_list, qty_col_header='Quantidade Total')

    wb.save(out_path)
    print(f'OK Excel salvo: {out_path}')


# ── Export DOCX ─────────────────────────────────────────────────
GRAY_DARK = '808080'
GRAY_MID  = 'AEAAAA'
CENTER = WD_ALIGN_PARAGRAPH.CENTER
COL_FULL = [1050, 1750, 1750, 1750, 2120, 2120]
COL_UNIT = [2000, 4270, 4270]

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.upper())
    tcPr.append(shd)

def set_col_width(cell, w_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def write_cell(cell, text, bold=False, align=CENTER, shade=None, size_pt=9):
    if shade:
        set_shading(cell, shade)
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(str(text) if text is not None else '')
    run.font.name = 'Calibri'
    run.font.bold = bold
    run.font.size = Pt(size_pt)

def apply_col_widths(table, col_widths):
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(col_widths):
                set_col_width(cell, col_widths[ci])

def add_memorial_table(doc, item):
    htype       = item['htype']
    code        = item['code']
    desc        = item['desc']
    rows        = item['rows']
    total       = item['total']
    unit        = item['unit']
    total_label = item['total_label']
    last_header = item['last_header']

    if htype == 'full':
        ncols = 6
        cw    = COL_FULL
        hdrs  = ['Foto', 'Comp. (m)', 'Altura (m)', 'Desconto (m2)', 'Subtotal', last_header]
    else:  # unit
        ncols = 3
        cw    = COL_UNIT
        hdrs  = ['Foto', 'Quantitativo', 'Unidade']

    nrows = 3 + len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: "Itens"
    r0 = table.rows[0]
    r0.cells[0].merge(r0.cells[-1])
    write_cell(r0.cells[0], 'Itens', bold=True, align=CENTER, shade=GRAY_DARK)
    set_col_width(r0.cells[0], sum(cw))

    # Row 1: Codigo | Descricao
    r1 = table.rows[1]
    write_cell(r1.cells[0], code, bold=True, shade=GRAY_DARK)
    set_col_width(r1.cells[0], cw[0])
    r1.cells[1].merge(r1.cells[-1])
    write_cell(r1.cells[1], desc, bold=True, shade=GRAY_DARK)

    # Row 2: Cabecalhos
    r2 = table.rows[2]
    for ci, hdr in enumerate(hdrs):
        write_cell(r2.cells[ci], hdr, bold=True, align=CENTER)
        set_col_width(r2.cells[ci], cw[ci])

    # Data rows
    for ri, r in enumerate(rows):
        tr = table.rows[3 + ri]
        if htype == 'full':
            write_cell(tr.cells[0], r.get('foto', ''), align=CENTER)
            write_cell(tr.cells[1], r.get('comp', ''), align=CENTER)
            write_cell(tr.cells[2], r.get('alt',  ''), align=CENTER)
            write_cell(tr.cells[3], r.get('desc', ''), align=CENTER)
            write_cell(tr.cells[4], r.get('sub',  ''), align=CENTER)
            write_cell(tr.cells[5], r.get('tot',  ''), align=CENTER)
        else:  # unit
            write_cell(tr.cells[0], r.get('foto', ''), align=CENTER)
            write_cell(tr.cells[1], r.get('qty',  ''), align=CENTER)
            write_cell(tr.cells[2], r.get('unit', ''), align=CENTER)
        for ci in range(ncols):
            set_col_width(tr.cells[ci], cw[ci])

    # Total row
    tr_tot = table.rows[-1]
    tr_tot.cells[0].merge(tr_tot.cells[-2])
    write_cell(tr_tot.cells[0], total_label, bold=True, align=CENTER, shade=GRAY_MID)
    val_str = f'{fmt(total)} {unit}' if unit and htype == 'unit' else fmt(total)
    write_cell(tr_tot.cells[-1], val_str, bold=True, align=CENTER, shade=GRAY_MID)

    apply_col_widths(table, cw)

def export_docx(consolidated, out_path):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin  = section.bottom_margin = Cm(2)

    for idx, (code, item) in enumerate(consolidated.items()):
        add_memorial_table(doc, item)
        if idx < len(consolidated) - 1:
            doc.add_paragraph('')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with open(out_path, 'wb') as f:
        f.write(buf.read())

    print(f'OK Memorial salvo: {out_path}')

if __name__ == '__main__':
    pass
