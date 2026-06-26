from docx import Document

docx_path = r"C:\Users\thiag\Downloads\- Teste execução - DANILO\RELATÓRIO FOTOGRÁFICO - VILA TIBERIO - LEVANTAMENTO PREVENTIVO(1).docx"

def get_cell_text(cell):
    return '\n'.join(p.text for p in cell.paragraphs).strip()

def raw_row(row):
    return [get_cell_text(c) for c in row.cells]

with open(docx_path, 'rb') as f:
    doc = Document(f)

print("Procurando o texto '13.12' em todas as tabelas, mesmo as ignoradas...")

count = 0
for t_idx, table in enumerate(doc.tables):
    found_in_table = False
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "13.12" in cell.text:
                found_in_table = True
                break
        if found_in_table:
            break
            
    if found_in_table:
        count += 1
        print(f"\n--- Tabela {t_idx} contém '13.12' ---")
        for r_idx, row in enumerate(table.rows):
            print(f"Row {r_idx}: {raw_row(row)}")
            
print(f"\nTotal de tabelas contendo '13.12': {count}")
